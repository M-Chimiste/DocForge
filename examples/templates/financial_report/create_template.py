"""Create Financial Report template and data files for DocForge.

Run: conda run -n docforge python examples/templates/financial_report/create_template.py
"""

import csv
import json
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor

TEMPLATE_DIR = Path(__file__).parent
DATA_DIR = TEMPLATE_DIR / "data"

DATA_DIR.mkdir(parents=True, exist_ok=True)


def _red(paragraph, text):
    """Add a red-colored run to a paragraph (exact #FF0000)."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def _red_cell(cell, text):
    """Set a table cell's text to red."""
    cell.text = ""
    p = cell.paragraphs[0]
    _red(p, text)


def create_template():
    """Create template.docx with red text markers for DocForge."""
    doc = Document()

    # Title
    doc.add_heading("Annual Financial Report", level=0)

    # --- Section 1: Company Overview ---
    doc.add_heading("Company Overview", level=1)
    p1 = doc.add_paragraph("Company: ")
    _red(p1, "Company Name")
    p2 = doc.add_paragraph("Fiscal Year: ")
    _red(p2, "Fiscal Year")
    p3 = doc.add_paragraph("Report Date: ")
    _red(p3, "Report Date")
    doc.add_paragraph(
        "This report provides a comprehensive overview of the company's financial "
        "performance for the fiscal year."
    )

    # --- Section 2: Revenue Summary ---
    doc.add_heading("Revenue Summary", level=1)
    doc.add_paragraph(
        "The following table summarizes quarterly revenue performance:"
    )

    # Skeleton table with header row + red sample data row
    table_rev = doc.add_table(rows=2, cols=3)
    table_rev.style = "Table Grid"

    # Header row (normal text)
    headers_rev = table_rev.rows[0].cells
    headers_rev[0].text = "Quarter"
    headers_rev[1].text = "Revenue"
    headers_rev[2].text = "Growth"

    # Sample data row (red text)
    sample_rev = table_rev.rows[1].cells
    _red_cell(sample_rev[0], "Q1 2025")
    _red_cell(sample_rev[1], "$125,000")
    _red_cell(sample_rev[2], "5.2%")

    # --- Section 3: Expense Analysis ---
    doc.add_heading("Expense Analysis", level=1)
    doc.add_paragraph(
        "The table below breaks down expenses by category:"
    )

    # Skeleton table with header row + red sample data row
    table_exp = doc.add_table(rows=2, cols=3)
    table_exp.style = "Table Grid"

    # Header row
    headers_exp = table_exp.rows[0].cells
    headers_exp[0].text = "Category"
    headers_exp[1].text = "Amount"
    headers_exp[2].text = "Percentage"

    # Sample data row (red text)
    sample_exp = table_exp.rows[1].cells
    _red_cell(sample_exp[0], "Salaries")
    _red_cell(sample_exp[1], "$450,000")
    _red_cell(sample_exp[2], "45.0%")

    # --- Section 4: Financial Outlook ---
    doc.add_heading("Financial Outlook", level=1)
    p_llm = doc.add_paragraph()
    _red(
        p_llm,
        "Write a brief financial outlook based on the revenue and expense data "
        "provided. Highlight key trends and recommendations.",
    )

    output_path = TEMPLATE_DIR / "template.docx"
    doc.save(str(output_path))
    print(f"Created {output_path}")


def create_data_files():
    """Create sample data files for the financial report template."""

    # --- revenue.xlsx ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Revenue"
    ws.append(["Quarter", "Revenue", "Growth"])
    ws.append(["Q1 2025", 125000, "5.2%"])
    ws.append(["Q2 2025", 148000, "18.4%"])
    ws.append(["Q3 2025", 132000, "-10.8%"])
    ws.append(["Q4 2025", 167000, "26.5%"])
    revenue_path = DATA_DIR / "revenue.xlsx"
    wb.save(str(revenue_path))
    print(f"Created {revenue_path}")

    # --- expenses.csv ---
    expenses_path = DATA_DIR / "expenses.csv"
    with open(expenses_path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Category", "Amount", "Percentage"])
        writer.writerow(["Salaries", 450000, "45.0%"])
        writer.writerow(["Marketing", 120000, "12.0%"])
        writer.writerow(["Operations", 180000, "18.0%"])
        writer.writerow(["Technology", 95000, "9.5%"])
        writer.writerow(["Rent & Utilities", 85000, "8.5%"])
        writer.writerow(["Miscellaneous", 70000, "7.0%"])
    print(f"Created {expenses_path}")

    # --- config.json ---
    config = {
        "company_name": "Acme Financial Corp",
        "fiscal_year": "FY2025",
        "report_date": "2025-12-31",
    }
    config_path = DATA_DIR / "config.json"
    with open(config_path, "w") as f:
        json.dump(config, f, indent=2)
    print(f"Created {config_path}")


def create_mapping():
    """Create mapping.json for the financial report template."""
    mapping = [
        {
            "marker_id": "marker-0",
            "data_source": "config.json",
            "field": "company_name",
        },
        {
            "marker_id": "marker-1",
            "data_source": "config.json",
            "field": "fiscal_year",
        },
        {
            "marker_id": "marker-2",
            "data_source": "config.json",
            "field": "report_date",
        },
        {
            "marker_id": "table-0",
            "data_source": "revenue.xlsx",
            "sheet": "Revenue",
        },
        {
            "marker_id": "table-1",
            "data_source": "expenses.csv",
        },
    ]
    mapping_path = TEMPLATE_DIR / "mapping.json"
    with open(mapping_path, "w") as f:
        json.dump(mapping, f, indent=2)
    print(f"Created {mapping_path}")


if __name__ == "__main__":
    create_template()
    create_data_files()
    create_mapping()
    print("\nFinancial Report template package created successfully!")
    print(
        f"\nTry: conda run -n docforge python -m cli analyze "
        f"--template {TEMPLATE_DIR / 'template.docx'}"
    )
