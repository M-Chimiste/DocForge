"""Create Compliance Assessment Report template and data files for DocForge.

Run: conda run -n docforge python examples/templates/compliance_document/create_template.py
"""

import csv
import json
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor

TEMPLATE_DIR = Path(__file__).parent
DATA_DIR = TEMPLATE_DIR / "data"

DATA_DIR.mkdir(parents=True, exist_ok=True)


def _red(paragraph, text):
    """Add a red-colored run to a paragraph (exact #FF0000)."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def _red_cell(cell, text):
    """Set a table cell's text to red."""
    cell.text = ""
    p = cell.paragraphs[0]
    _red(p, text)


def create_template():
    """Create template.docx with red text markers for DocForge."""
    doc = Document()

    # Title
    doc.add_heading("Compliance Assessment Report", level=0)

    # --- Section 1: Assessment Overview ---
    doc.add_heading("Assessment Overview", level=1)
    p1 = doc.add_paragraph("Organization: ")
    _red(p1, "Organization Name")
    p2 = doc.add_paragraph("Assessment Date: ")
    _red(p2, "Assessment Date")
    p3 = doc.add_paragraph("Lead Assessor: ")
    _red(p3, "Assessor")
    doc.add_paragraph(
        "This document presents the results of the compliance assessment "
        "conducted against the applicable regulatory framework. It covers "
        "control effectiveness, identified findings, and remediation "
        "recommendations."
    )

    # --- Section 2: Control Checklist ---
    doc.add_heading("Control Checklist", level=1)
    doc.add_paragraph(
        "The following table lists all assessed controls and their current status:"
    )

    # Skeleton table with header row + red sample data row
    table_ctrl = doc.add_table(rows=2, cols=4)
    table_ctrl.style = "Table Grid"

    # Header row
    ctrl_headers = table_ctrl.rows[0].cells
    ctrl_headers[0].text = "Control ID"
    ctrl_headers[1].text = "Description"
    ctrl_headers[2].text = "Status"
    ctrl_headers[3].text = "Evidence"

    # Sample data row (red text)
    ctrl_sample = table_ctrl.rows[1].cells
    _red_cell(ctrl_sample[0], "AC-01")
    _red_cell(ctrl_sample[1], "Access control policy")
    _red_cell(ctrl_sample[2], "Compliant")
    _red_cell(ctrl_sample[3], "Policy document v2.3")

    # --- Section 3: Findings ---
    doc.add_heading("Findings", level=1)
    doc.add_paragraph(
        "The table below details all findings identified during the assessment:"
    )

    # Skeleton table with header row + red sample data row
    table_find = doc.add_table(rows=2, cols=4)
    table_find.style = "Table Grid"

    # Header row
    find_headers = table_find.rows[0].cells
    find_headers[0].text = "Finding ID"
    find_headers[1].text = "Severity"
    find_headers[2].text = "Description"
    find_headers[3].text = "Affected Control"

    # Sample data row (red text)
    find_sample = table_find.rows[1].cells
    _red_cell(find_sample[0], "F-001")
    _red_cell(find_sample[1], "Critical")
    _red_cell(find_sample[2], "Missing MFA enforcement")
    _red_cell(find_sample[3], "AC-03")

    # --- Section 4: Remediation Plan ---
    doc.add_heading("Remediation Plan", level=1)
    p_llm = doc.add_paragraph()
    _red(
        p_llm,
        "Based on the control checklist and findings, provide a remediation plan "
        "with prioritized recommendations. Focus on critical and high-severity "
        "findings first.",
    )

    output_path = TEMPLATE_DIR / "template.docx"
    doc.save(str(output_path))
    print(f"Created {output_path}")


def create_data_files():
    """Create sample data files for the compliance template."""

    # --- controls.xlsx ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Controls"
    ws.append(["Control ID", "Description", "Status", "Evidence"])
    ws.append(["AC-01", "Access control policy", "Compliant", "Policy document v2.3"])
    ws.append(
        ["AC-02", "User provisioning process", "Compliant", "HR workflow records"]
    )
    ws.append(
        [
            "AC-03",
            "Multi-factor authentication",
            "Non-Compliant",
            "MFA audit log gaps",
        ]
    )
    ws.append(
        ["AC-04", "Privileged access management", "Compliant", "PAM tool audit trail"]
    )
    ws.append(
        [
            "CM-01",
            "Configuration management baseline",
            "Compliant",
            "Baseline config docs",
        ]
    )
    ws.append(
        [
            "CM-02",
            "Change management process",
            "Partially Compliant",
            "Change tickets reviewed",
        ]
    )
    ws.append(
        [
            "IR-01",
            "Incident response plan",
            "Compliant",
            "IR plan v4.1 and drill records",
        ]
    )
    ws.append(
        [
            "IR-02",
            "Incident escalation procedures",
            "Non-Compliant",
            "No documented escalation matrix",
        ]
    )
    ws.append(
        [
            "AU-01",
            "Audit logging enabled",
            "Compliant",
            "SIEM dashboard screenshots",
        ]
    )
    ws.append(
        [
            "AU-02",
            "Log retention policy",
            "Partially Compliant",
            "90-day retention vs 365-day requirement",
        ]
    )
    controls_path = DATA_DIR / "controls.xlsx"
    wb.save(str(controls_path))
    print(f"Created {controls_path}")

    # --- findings.csv ---
    findings_path = DATA_DIR / "findings.csv"
    with open(findings_path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Finding ID", "Severity", "Description", "Affected Control"])
        writer.writerow(
            [
                "F-001",
                "Critical",
                "MFA not enforced for administrative accounts",
                "AC-03",
            ]
        )
        writer.writerow(
            [
                "F-002",
                "High",
                "No documented incident escalation matrix",
                "IR-02",
            ]
        )
        writer.writerow(
            [
                "F-003",
                "Medium",
                "Change management approvals missing for 3 of 20 sampled changes",
                "CM-02",
            ]
        )
        writer.writerow(
            [
                "F-004",
                "Medium",
                "Audit log retention set to 90 days instead of required 365 days",
                "AU-02",
            ]
        )
        writer.writerow(
            [
                "F-005",
                "Low",
                "Access review documentation lacks reviewer signatures",
                "AC-02",
            ]
        )
    print(f"Created {findings_path}")

    # --- assessment.json ---
    assessment = {
        "organization_name": "GlobalTech Solutions Inc.",
        "assessment_date": "2025-10-15",
        "assessor": "Maria Chen, CISA",
    }
    assessment_path = DATA_DIR / "assessment.json"
    with open(assessment_path, "w") as f:
        json.dump(assessment, f, indent=2)
    print(f"Created {assessment_path}")


def create_mapping():
    """Create mapping.json for the compliance template."""
    mapping = [
        {
            "marker_id": "marker-0",
            "data_source": "assessment.json",
            "field": "organization_name",
        },
        {
            "marker_id": "marker-1",
            "data_source": "assessment.json",
            "field": "assessment_date",
        },
        {
            "marker_id": "marker-2",
            "data_source": "assessment.json",
            "field": "assessor",
        },
        {
            "marker_id": "table-0",
            "data_source": "controls.xlsx",
            "sheet": "Controls",
        },
        {
            "marker_id": "table-1",
            "data_source": "findings.csv",
        },
    ]
    mapping_path = TEMPLATE_DIR / "mapping.json"
    with open(mapping_path, "w") as f:
        json.dump(mapping, f, indent=2)
    print(f"Created {mapping_path}")


if __name__ == "__main__":
    create_template()
    create_data_files()
    create_mapping()
    print("\nCompliance Assessment Report template package created successfully!")
    print(
        f"\nTry: conda run -n docforge python -m cli analyze "
        f"--template {TEMPLATE_DIR / 'template.docx'}"
    )
