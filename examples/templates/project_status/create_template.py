"""Create Project Status Report template and data files for DocForge.

Run: conda run -n docforge python examples/templates/project_status/create_template.py
"""

import csv
import json
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, RGBColor

TEMPLATE_DIR = Path(__file__).parent
DATA_DIR = TEMPLATE_DIR / "data"

DATA_DIR.mkdir(parents=True, exist_ok=True)


def _red(paragraph, text):
    """Add a red-colored run to a paragraph (exact #FF0000)."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def _red_cell(cell, text):
    """Set a table cell's text to red."""
    cell.text = ""
    p = cell.paragraphs[0]
    _red(p, text)


def create_template():
    """Create template.docx with red text markers for DocForge."""
    doc = Document()

    # Title
    doc.add_heading("Project Status Report", level=0)

    # --- Section 1: Project Details ---
    doc.add_heading("Project Details", level=1)
    p1 = doc.add_paragraph("Project: ")
    _red(p1, "Project Name")
    p2 = doc.add_paragraph("Sprint: ")
    _red(p2, "Sprint Number")
    p3 = doc.add_paragraph("Report Date: ")
    _red(p3, "Report Date")
    doc.add_paragraph(
        "This report provides the current status of the project including "
        "sprint performance metrics and identified risks."
    )

    # --- Section 2: Sprint Metrics ---
    doc.add_heading("Sprint Metrics", level=1)
    doc.add_paragraph(
        "The following table tracks sprint-level performance across the project:"
    )

    # Skeleton table with header row + red sample data row
    table_sprint = doc.add_table(rows=2, cols=4)
    table_sprint.style = "Table Grid"

    # Header row
    headers = table_sprint.rows[0].cells
    headers[0].text = "Sprint"
    headers[1].text = "Planned Points"
    headers[2].text = "Completed Points"
    headers[3].text = "Velocity"

    # Sample data row (red text)
    sample = table_sprint.rows[1].cells
    _red_cell(sample[0], "Sprint 1")
    _red_cell(sample[1], "40")
    _red_cell(sample[2], "36")
    _red_cell(sample[3], "90%")

    # --- Section 3: Risk Register ---
    doc.add_heading("Risk Register", level=1)
    doc.add_paragraph(
        "The table below captures identified risks and their mitigation strategies:"
    )

    # Skeleton table with header row + red sample data row
    table_risk = doc.add_table(rows=2, cols=4)
    table_risk.style = "Table Grid"

    # Header row
    risk_headers = table_risk.rows[0].cells
    risk_headers[0].text = "Risk ID"
    risk_headers[1].text = "Description"
    risk_headers[2].text = "Severity"
    risk_headers[3].text = "Mitigation"

    # Sample data row (red text)
    risk_sample = table_risk.rows[1].cells
    _red_cell(risk_sample[0], "R-001")
    _red_cell(risk_sample[1], "Resource availability")
    _red_cell(risk_sample[2], "High")
    _red_cell(risk_sample[3], "Cross-train team members")

    # --- Section 4: Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    p_llm = doc.add_paragraph()
    _red(
        p_llm,
        "Provide an executive summary of the project status based on the sprint "
        "metrics and risk register. Include recommendations for the next sprint.",
    )

    output_path = TEMPLATE_DIR / "template.docx"
    doc.save(str(output_path))
    print(f"Created {output_path}")


def create_data_files():
    """Create sample data files for the project status template."""

    # --- sprints.csv ---
    sprints_path = DATA_DIR / "sprints.csv"
    with open(sprints_path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Sprint", "Planned Points", "Completed Points", "Velocity"])
        writer.writerow(["Sprint 1", 40, 36, "90%"])
        writer.writerow(["Sprint 2", 45, 42, "93%"])
        writer.writerow(["Sprint 3", 50, 38, "76%"])
        writer.writerow(["Sprint 4", 42, 41, "98%"])
        writer.writerow(["Sprint 5", 48, 44, "92%"])
    print(f"Created {sprints_path}")

    # --- risks.yaml ---
    risks = [
        {
            "risk_id": "R-001",
            "description": "Resource availability due to concurrent projects",
            "severity": "High",
            "mitigation": "Cross-train team members and identify backup resources",
        },
        {
            "risk_id": "R-002",
            "description": "Third-party API integration delays",
            "severity": "Medium",
            "mitigation": "Build mock services for parallel development",
        },
        {
            "risk_id": "R-003",
            "description": "Scope creep from stakeholder requests",
            "severity": "High",
            "mitigation": "Enforce change control process and sprint boundaries",
        },
        {
            "risk_id": "R-004",
            "description": "Performance degradation under load",
            "severity": "Medium",
            "mitigation": "Schedule load testing in Sprint 6",
        },
        {
            "risk_id": "R-005",
            "description": "Key team member leaving the project",
            "severity": "Low",
            "mitigation": "Document knowledge and maintain updated onboarding guide",
        },
    ]
    risks_path = DATA_DIR / "risks.yaml"
    with open(risks_path, "w") as f:
        yaml.dump(risks, f, default_flow_style=False, sort_keys=False)
    print(f"Created {risks_path}")

    # --- project.json ---
    project = {
        "project_name": "Phoenix Platform Migration",
        "sprint_number": "Sprint 5",
        "report_date": "2025-11-14",
    }
    project_path = DATA_DIR / "project.json"
    with open(project_path, "w") as f:
        json.dump(project, f, indent=2)
    print(f"Created {project_path}")


def create_mapping():
    """Create mapping.json for the project status template."""
    mapping = [
        {
            "marker_id": "marker-0",
            "data_source": "project.json",
            "field": "project_name",
        },
        {
            "marker_id": "marker-1",
            "data_source": "project.json",
            "field": "sprint_number",
        },
        {
            "marker_id": "marker-2",
            "data_source": "project.json",
            "field": "report_date",
        },
        {
            "marker_id": "table-0",
            "data_source": "sprints.csv",
        },
        {
            "marker_id": "table-1",
            "data_source": "risks.yaml",
        },
    ]
    mapping_path = TEMPLATE_DIR / "mapping.json"
    with open(mapping_path, "w") as f:
        json.dump(mapping, f, indent=2)
    print(f"Created {mapping_path}")


if __name__ == "__main__":
    create_template()
    create_data_files()
    create_mapping()
    print("\nProject Status Report template package created successfully!")
    print(
        f"\nTry: conda run -n docforge python -m cli analyze "
        f"--template {TEMPLATE_DIR / 'template.docx'}"
    )
