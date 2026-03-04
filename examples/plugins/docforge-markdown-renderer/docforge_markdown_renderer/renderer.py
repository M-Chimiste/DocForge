"""Markdown renderer plugin for DocForge.

Handles LLM_PROMPT markers whose text starts with ``[markdown]``.  The
renderer fetches the resolved text from the data store, parses basic
Markdown formatting (bold and italic), and writes the result into the
document paragraph with appropriate run-level formatting.
"""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx import Document
from docx.shared import Pt

from core.data_loader import DataStore
from core.models import MappingEntry, MarkerType, RenderResult, TemplateMarker
from renderers.base import BaseRenderer
from utils.formatting import clear_red_color


# ------------------------------------------------------------------
# Lightweight Markdown token model
# ------------------------------------------------------------------

@dataclass
class _TextFragment:
    """A piece of text with optional bold / italic flags."""

    text: str
    bold: bool = False
    italic: bool = False


_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*"   # ***bold italic***
    r"|\*\*(.+?)\*\*"       # **bold**
    r"|\*(.+?)\*)"          # *italic*
)


def _parse_inline(text: str) -> list[_TextFragment]:
    """Split *text* into fragments honouring ``**bold**`` and ``*italic*``."""
    fragments: list[_TextFragment] = []
    last = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > last:
            fragments.append(_TextFragment(text=text[last : m.start()]))
        if m.group(2):
            fragments.append(_TextFragment(text=m.group(2), bold=True, italic=True))
        elif m.group(3):
            fragments.append(_TextFragment(text=m.group(3), bold=True))
        elif m.group(4):
            fragments.append(_TextFragment(text=m.group(4), italic=True))
        last = m.end()
    if last < len(text):
        fragments.append(_TextFragment(text=text[last:]))
    return fragments


def _strip_heading(line: str) -> tuple[str, int]:
    """Return ``(text, level)`` for Markdown headings (``# …`` through ``### …``).

    Non-heading lines return ``(line, 0)``.
    """
    stripped = line.lstrip()
    level = 0
    while level < len(stripped) and stripped[level] == "#":
        level += 1
    if level > 0 and level < len(stripped) and stripped[level] == " ":
        return stripped[level + 1 :], min(level, 3)
    return line, 0


# ------------------------------------------------------------------
# Renderer
# ------------------------------------------------------------------

class MarkdownRenderer(BaseRenderer):
    """Render ``[markdown]``-prefixed LLM prompt markers.

    The marker text is expected to look like::

        [markdown] Some instruction for the LLM

    The renderer resolves the actual content from the mapped data source,
    converts basic Markdown (bold / italic / headings) into docx run
    formatting, and replaces the original red-text paragraph.
    """

    def can_handle(self, marker: TemplateMarker) -> bool:
        return (
            marker.marker_type == MarkerType.LLM_PROMPT
            and marker.text.strip().lower().startswith("[markdown]")
        )

    def render(
        self,
        marker: TemplateMarker,
        data: DataStore,
        document: Document,
        mapping: MappingEntry,
    ) -> RenderResult:
        # --- resolve content ------------------------------------------------
        content = self._resolve_content(data, mapping)
        if content is None:
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error=f"Could not resolve markdown content from {mapping.data_source}",
            )

        # --- locate the paragraph -------------------------------------------
        if marker.paragraph_index < 0 or marker.paragraph_index >= len(document.paragraphs):
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error=f"Paragraph index {marker.paragraph_index} out of range",
            )

        paragraph = document.paragraphs[marker.paragraph_index]

        # --- clear existing runs (the red-text content) ---------------------
        for run in paragraph.runs:
            run.text = ""
            clear_red_color(run)

        # --- write markdown-formatted runs ----------------------------------
        for line in content.splitlines():
            text, heading_level = _strip_heading(line)
            fragments = _parse_inline(text)
            for frag in fragments:
                run = paragraph.add_run(frag.text)
                run.bold = frag.bold or heading_level > 0
                run.italic = frag.italic
                if heading_level > 0:
                    run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(heading_level, 12))

        return RenderResult(
            marker_id=marker.id, success=True, rendered_by="markdown"
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _resolve_content(data: DataStore, mapping: MappingEntry) -> str | None:
        """Get markdown text from the data store."""
        # Try text content first (unstructured sources)
        text = data.get_text(mapping.data_source)
        if text:
            return text

        # Fall back to a DataFrame field
        df = data.get_dataframe(mapping.data_source, sheet=mapping.sheet)
        if df is not None and mapping.field and mapping.field in df.columns:
            return str(df.iloc[0][mapping.field])

        return None
