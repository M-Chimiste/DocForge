"""Tests for the MarkdownRenderer plugin."""

from __future__ import annotations

import pytest
from docx import Document

from core.data_loader import DataStore
from core.models import MappingEntry, MarkerType, RenderResult, TemplateMarker
from extractors.base import ExtractedData
from docforge_markdown_renderer.renderer import (
    MarkdownRenderer,
    _parse_inline,
    _strip_heading,
)

from pathlib import Path


# ------------------------------------------------------------------
# Unit tests for inline parser
# ------------------------------------------------------------------


class TestParseInline:
    def test_plain_text(self):
        frags = _parse_inline("hello world")
        assert len(frags) == 1
        assert frags[0].text == "hello world"
        assert not frags[0].bold
        assert not frags[0].italic

    def test_bold(self):
        frags = _parse_inline("some **bold** text")
        assert len(frags) == 3
        assert frags[1].text == "bold"
        assert frags[1].bold

    def test_italic(self):
        frags = _parse_inline("some *italic* text")
        assert len(frags) == 3
        assert frags[1].text == "italic"
        assert frags[1].italic

    def test_bold_italic(self):
        frags = _parse_inline("***both***")
        assert len(frags) == 1
        assert frags[0].bold
        assert frags[0].italic


class TestStripHeading:
    def test_h1(self):
        text, level = _strip_heading("# Title")
        assert text == "Title"
        assert level == 1

    def test_h2(self):
        text, level = _strip_heading("## Subtitle")
        assert text == "Subtitle"
        assert level == 2

    def test_plain(self):
        text, level = _strip_heading("Just text")
        assert text == "Just text"
        assert level == 0


# ------------------------------------------------------------------
# Integration-style tests
# ------------------------------------------------------------------


class TestMarkdownRendererCanHandle:
    def setup_method(self):
        self.renderer = MarkdownRenderer()

    def test_matches_markdown_prefix(self):
        marker = TemplateMarker(
            id="m1",
            text="[markdown] Write a summary",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        assert self.renderer.can_handle(marker)

    def test_rejects_non_llm_prompt(self):
        marker = TemplateMarker(
            id="m2",
            text="[markdown] something",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[0],
        )
        assert not self.renderer.can_handle(marker)

    def test_rejects_without_prefix(self):
        marker = TemplateMarker(
            id="m3",
            text="Write a summary",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        assert not self.renderer.can_handle(marker)


class TestMarkdownRendererRender:
    def setup_method(self):
        self.renderer = MarkdownRenderer()

    def _make_doc_with_red_paragraph(self, text: str = "placeholder") -> Document:
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run(text)
        return doc

    def _make_data_store(self, text_content: str) -> DataStore:
        store = DataStore()
        ed = ExtractedData(source_path=Path("test.md"), text_content=text_content)
        store.add("test.md", ed)
        return store

    def test_render_replaces_text(self):
        doc = self._make_doc_with_red_paragraph()
        store = self._make_data_store("Hello **world**")
        marker = TemplateMarker(
            id="m1",
            text="[markdown] Write greeting",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        mapping = MappingEntry(marker_id="m1", data_source="test.md")
        result = self.renderer.render(marker, store, doc, mapping)

        assert result.success
        assert result.rendered_by == "markdown"
        # The paragraph should now contain the rendered content
        full_text = doc.paragraphs[0].text
        assert "Hello" in full_text
        assert "world" in full_text

    def test_render_fails_on_missing_data(self):
        doc = self._make_doc_with_red_paragraph()
        store = DataStore()
        marker = TemplateMarker(
            id="m1",
            text="[markdown] Write greeting",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        mapping = MappingEntry(marker_id="m1", data_source="missing.md")
        result = self.renderer.render(marker, store, doc, mapping)

        assert not result.success
        assert "Could not resolve" in result.error
