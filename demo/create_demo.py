"""Create demo template and data files for DocForge.

Run: conda run -n docforge python demo/create_demo.py
"""

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import RGBColor

DEMO_DIR = Path(__file__).parent
TEMPLATES_DIR = DEMO_DIR / "templates"
DATA_DIR = DEMO_DIR / "data"

TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)


def _red(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def create_quarterly_report():
    doc = Document()

    # Title section
    doc.add_heading("Quarterly Report", level=0)

    # Project Info
    doc.add_heading("Project Information", level=1)
    p1 = doc.add_paragraph("Project: ")
    _red(p1, "Project Name")
    p2 = doc.add_paragraph("Prepared by: ")
    _red(p2, "Author")
    p2.add_run(" on ")
    _red(p2, "Report Date")

    # Financial Summary
    doc.add_heading("Financial Summary", level=1)
    doc.add_paragraph("The following table shows quarterly revenue performance:")
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = "Table Grid"
    headers = table1.rows[0].cells
    headers[0].text = "Quarter"
    headers[1].text = "Revenue"
    headers[2].text = "Growth"

    # Project Status
    doc.add_heading("Project Status", level=1)
    doc.add_paragraph("Current status of active projects:")
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Table Grid"
    headers2 = table2.rows[0].cells
    headers2[0].text = "Project"
    headers2[1].text = "Status"
    headers2[2].text = "Progress"
    headers2[3].text = "Lead"

    # Executive Summary (LLM prompt marker)
    doc.add_heading("Executive Summary", level=1)
    p3 = doc.add_paragraph()
    _red(
        p3,
        "Write a brief executive summary covering the key findings from the quarterly revenue data and project status updates",
    )

    doc.save(str(TEMPLATES_DIR / "quarterly_report.docx"))
    print("Created quarterly_report.docx")


def create_demo_data():
    # Excel with revenue data
    with pd.ExcelWriter(str(DATA_DIR / "quarterly_metrics.xlsx"), engine="openpyxl") as writer:
        df = pd.DataFrame(
            {
                "Quarter": ["Q1 2025", "Q2 2025", "Q3 2025", "Q4 2025"],
                "Revenue": [125000, 148000, 132000, 167000],
                "Growth": ["5.2%", "18.4%", "-10.8%", "26.5%"],
            }
        )
        df.to_excel(writer, sheet_name="Revenue", index=False)
    print("Created quarterly_metrics.xlsx")

    # CSV with project status
    df_csv = pd.DataFrame(
        {
            "Project": ["Atlas", "Beacon", "Cascade", "Delta"],
            "Status": ["Active", "Planning", "Complete", "Active"],
            "Progress": [78, 15, 100, 42],
            "Lead": ["Alice Chen", "Bob Martinez", "Carol Singh", "Dave Kim"],
        }
    )
    df_csv.to_csv(str(DATA_DIR / "project_status.csv"), index=False)
    print("Created project_status.csv")

    # JSON config
    config = {
        "project": {
            "name": "Acme Corp Q4 Report",
            "date": "2025-12-31",
            "version": "1.0",
        },
        "settings": {
            "author": "Jane Doe",
            "department": "Finance",
        },
    }
    with open(DATA_DIR / "config.json", "w") as f:
        json.dump(config, f, indent=2)
    print("Created config.json")


def create_mapping():
    mapping = [
        {"marker_id": "marker-0", "data_source": "config.json", "field": "name", "path": "project"},
        {"marker_id": "marker-1", "data_source": "config.json", "field": "author", "path": "settings"},
        {"marker_id": "marker-2", "data_source": "config.json", "field": "date", "path": "project"},
        {"marker_id": "table-0", "data_source": "quarterly_metrics.xlsx", "sheet": "Revenue"},
        {"marker_id": "table-1", "data_source": "project_status.csv"},
    ]
    with open(DATA_DIR / "mapping.json", "w") as f:
        json.dump(mapping, f, indent=2)
    print("Created mapping.json")


if __name__ == "__main__":
    create_quarterly_report()
    create_demo_data()
    create_mapping()
    print("\nDemo files created successfully!")
    print(f"\nTry: conda run -n docforge python -m cli analyze --template {TEMPLATES_DIR / 'quarterly_report.docx'}")
