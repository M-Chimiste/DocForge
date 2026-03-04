"""Word (.docx) content extractor using markitdown + python-docx for tables."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from markitdown import MarkItDown

from extractors.base import BaseExtractor, ExtractedData, ExtractionConfig


class DocxContentExtractor(BaseExtractor):
    """Extracts content from .docx files used as data sources (not templates)."""

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def extract(self, file_path: Path, config: ExtractionConfig | None = None) -> ExtractedData:
        from docx import Document

        config = config or ExtractionConfig()

        # Use markitdown for rich markdown text extraction
        md = MarkItDown(enable_plugins=False)
        conversion = md.convert(str(file_path))
        full_text = conversion.text_content or ""

        result = ExtractedData(source_path=file_path, text_content=full_text)
        result.dataframes["default"] = pd.DataFrame({"content": [full_text]})

        # Use python-docx for structured table extraction
        doc = Document(str(file_path))
        for i, table in enumerate(doc.tables):
            if not table.rows:
                continue
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            result.dataframes[f"table_{i}"] = pd.DataFrame(rows, columns=headers)

        result.metadata["paragraph_count"] = len([p for p in doc.paragraphs if p.text.strip()])
        result.metadata["table_count"] = len(doc.tables)
        return result
