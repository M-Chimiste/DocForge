"""Low-level python-docx and lxml utilities for DocForge."""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from utils.formatting import clear_red_color


def get_heading_level(paragraph: Paragraph) -> int | None:
    """Return heading level (1-9) if paragraph uses a Heading style, else None."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return None
    return None


def copy_run_format(source_run: Run, target_run: Run) -> None:
    """Copy font formatting from source to target run, excluding color.

    This is used when replacing red text with mapped values — the replacement
    should carry the surrounding formatting but NOT the red color.
    """
    sf = source_run.font
    tf = target_run.font
    tf.bold = sf.bold
    tf.italic = sf.italic
    tf.underline = sf.underline
    tf.size = sf.size
    tf.name = sf.name
    # Explicitly do NOT copy color — we want to remove the red


def find_adjacent_non_red_run(paragraph: Paragraph, red_run_indices: list[int]) -> Run | None:
    """Find the nearest non-red run adjacent to the red runs for format copying."""
    runs = paragraph.runs

    # Try run before the first red run
    if red_run_indices and red_run_indices[0] > 0:
        return runs[red_run_indices[0] - 1]

    # Try run after the last red run
    if red_run_indices and red_run_indices[-1] < len(runs) - 1:
        return runs[red_run_indices[-1] + 1]

    return None


def inject_text_at_marker(
    marker_paragraph_index: int,
    marker_run_indices: list[int],
    document: Document,
    text: str,
) -> bool:
    """Replace red text runs with the given text, adding extra paragraphs for newlines.

    Returns True on success, False if paragraph index is out of range.
    """
    if marker_paragraph_index < 0 or marker_paragraph_index >= len(document.paragraphs):
        return False

    paragraph = document.paragraphs[marker_paragraph_index]
    runs = paragraph.runs

    format_source = find_adjacent_non_red_run(paragraph, marker_run_indices)

    text_parts = text.strip().split("\n")

    for i, run_idx in enumerate(marker_run_indices):
        if run_idx >= len(runs):
            continue
        run = runs[run_idx]
        if i == 0:
            run.text = text_parts[0] if text_parts else ""
        else:
            run.text = ""
        clear_red_color(run)
        if format_source:
            copy_run_format(format_source, run)

    if len(text_parts) > 1:
        current_para = paragraph
        for extra_text in text_parts[1:]:
            if extra_text.strip():
                new_para = _insert_paragraph_after(current_para, extra_text)
                if format_source and new_para.runs:
                    copy_run_format(format_source, new_para.runs[0])
                current_para = new_para

    return True


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Insert a new paragraph after the given paragraph."""
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para
