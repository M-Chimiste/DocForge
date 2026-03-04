"""Tests for the TipTap JSON to .docx export converter."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from core.docx_to_editor import DocxToEditorConverter
from core.editor_models import EditorDocument, EditorDocumentMeta, EditorMark, EditorNode
from core.editor_to_docx import EditorToDocxConverter
from core.models import GenerationReport, TemplateAnalysis
from utils.red_text import is_red_run


def _make_editor_doc(content_nodes: list[EditorNode]) -> EditorDocument:
    return EditorDocument(
        content=EditorNode(type="doc", content=content_nodes),
        meta=EditorDocumentMeta(
            generation_run_id=1,
            project_id=1,
        ),
    )


class TestBasicExport:
    def test_empty_document(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc([])
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert len(doc.paragraphs) == 0

    def test_single_paragraph(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="paragraph",
                    content=[EditorNode(type="text", text="Hello world")],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Hello world"

    def test_heading_export(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="heading",
                    attrs={"level": 1},
                    content=[EditorNode(type="text", text="Title")],
                ),
                EditorNode(
                    type="heading",
                    attrs={"level": 2},
                    content=[EditorNode(type="text", text="Subtitle")],
                ),
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert doc.paragraphs[0].text == "Title"
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[1].text == "Subtitle"
        assert doc.paragraphs[1].style.name == "Heading 2"


class TestFormattingExport:
    def test_bold_text(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="paragraph",
                    content=[
                        EditorNode(
                            type="text",
                            text="Bold",
                            marks=[EditorMark(type="bold")],
                        )
                    ],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert doc.paragraphs[0].runs[0].font.bold is True

    def test_italic_text(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="paragraph",
                    content=[
                        EditorNode(
                            type="text",
                            text="Italic",
                            marks=[EditorMark(type="italic")],
                        )
                    ],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert doc.paragraphs[0].runs[0].font.italic is True

    def test_text_style_font_size(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="paragraph",
                    content=[
                        EditorNode(
                            type="text",
                            text="Big text",
                            marks=[EditorMark(type="textStyle", attrs={"fontSize": "16.0pt"})],
                        )
                    ],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert doc.paragraphs[0].runs[0].font.size == Pt(16)


class TestUnresolvedExport:
    def test_unresolved_marker_exports_as_red(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="paragraph",
                    content=[
                        EditorNode(type="text", text="Before "),
                        EditorNode(
                            type="docforgeUnresolved",
                            attrs={"markerId": "m-0", "originalText": "Missing Data"},
                        ),
                        EditorNode(type="text", text=" after"),
                    ],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        para = doc.paragraphs[0]
        # Should have 3 runs: normal, red, normal
        assert len(para.runs) == 3
        assert para.runs[1].text == "Missing Data"
        assert is_red_run(para.runs[1]._element)


class TestTableExport:
    def test_basic_table(self, tmp_path):
        template = tmp_path / "template.docx"
        Document().save(str(template))

        editor_doc = _make_editor_doc(
            [
                EditorNode(
                    type="table",
                    attrs={"docforgeTableId": "table-0"},
                    content=[
                        EditorNode(
                            type="tableRow",
                            content=[
                                EditorNode(
                                    type="tableHeader",
                                    content=[
                                        EditorNode(
                                            type="paragraph",
                                            content=[EditorNode(type="text", text="Name")],
                                        )
                                    ],
                                ),
                                EditorNode(
                                    type="tableHeader",
                                    content=[
                                        EditorNode(
                                            type="paragraph",
                                            content=[EditorNode(type="text", text="Value")],
                                        )
                                    ],
                                ),
                            ],
                        ),
                        EditorNode(
                            type="tableRow",
                            content=[
                                EditorNode(
                                    type="tableCell",
                                    content=[
                                        EditorNode(
                                            type="paragraph",
                                            content=[EditorNode(type="text", text="Revenue")],
                                        )
                                    ],
                                ),
                                EditorNode(
                                    type="tableCell",
                                    content=[
                                        EditorNode(
                                            type="paragraph",
                                            content=[EditorNode(type="text", text="100000")],
                                        )
                                    ],
                                ),
                            ],
                        ),
                    ],
                )
            ]
        )
        converter = EditorToDocxConverter(editor_doc, template)
        doc = converter.convert()

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert table.rows[0].cells[0].text == "Name"
        assert table.rows[1].cells[0].text == "Revenue"
        assert table.rows[1].cells[1].text == "100000"


class TestRoundTrip:
    def test_paragraph_round_trip(self, tmp_path):
        """Create doc → convert to editor → export back → verify content."""
        template = tmp_path / "template.docx"
        original = Document()
        original.add_heading("Title", level=1)
        original.add_paragraph("Hello world")
        original.save(str(template))

        # Convert to editor format
        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )
        converter = DocxToEditorConverter(original, report, analysis, [])
        editor_doc = converter.convert()

        # Export back
        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        # Verify content preserved
        assert exported.paragraphs[0].text == "Title"
        assert exported.paragraphs[0].style.name == "Heading 1"
        assert exported.paragraphs[1].text == "Hello world"

    def test_table_round_trip(self, tmp_path):
        """Table content survives round-trip conversion."""
        template = tmp_path / "template.docx"
        original = Document()
        table = original.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Alpha"
        table.rows[1].cells[1].text = "100"
        original.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )
        converter = DocxToEditorConverter(original, report, analysis, [])
        editor_doc = converter.convert()

        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        assert len(exported.tables) == 1
        assert exported.tables[0].rows[0].cells[0].text == "Name"
        assert exported.tables[0].rows[1].cells[0].text == "Alpha"
