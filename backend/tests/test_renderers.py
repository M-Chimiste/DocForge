"""Tests for placeholder and table renderers."""

from pathlib import Path

import pytest
from docx import Document

from core.data_loader import DataStore, load_data_sources
from core.models import MappingEntry, MarkerType, TemplateMarker
from core.template_parser import parse_template
from renderers.placeholder_renderer import PlaceholderRenderer
from renderers.table_renderer import TableRenderer, render_table_direct

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture
def simple_placeholder_path():
    return FIXTURES_DIR / "templates" / "simple_placeholder.docx"


@pytest.fixture
def simple_table_path():
    return FIXTURES_DIR / "templates" / "simple_table.docx"


@pytest.fixture
def table_with_sample_data_path():
    return FIXTURES_DIR / "templates" / "table_with_sample_data.docx"


def _make_data_store_with_csv():
    """Create a DataStore with project_status.csv loaded."""
    return load_data_sources([FIXTURES_DIR / "data" / "project_status.csv"])


def _make_data_store_with_json():
    """Create a DataStore with config.json loaded."""
    return load_data_sources([FIXTURES_DIR / "data" / "config.json"])


class TestPlaceholderRenderer:
    def test_can_handle(self):
        renderer = PlaceholderRenderer()
        placeholder = TemplateMarker(
            id="m-0",
            text="Name",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[0],
        )
        llm_prompt = TemplateMarker(
            id="m-1",
            text="Summarize",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        assert renderer.can_handle(placeholder) is True
        assert renderer.can_handle(llm_prompt) is False

    def test_substitutes_placeholder(self, simple_placeholder_path, tmp_path):
        analysis = parse_template(simple_placeholder_path)
        doc = Document(str(simple_placeholder_path))
        store = _make_data_store_with_json()

        marker = next(
            m for m in analysis.markers if m.marker_type == MarkerType.VARIABLE_PLACEHOLDER
        )
        mapping = MappingEntry(
            marker_id=marker.id,
            data_source="config.json",
            path="project.name",
        )

        renderer = PlaceholderRenderer()
        result = renderer.render(marker, store, doc, mapping)
        assert result.success is True

        # Check the paragraph text was replaced
        para = doc.paragraphs[marker.paragraph_index]
        assert "DocForge Demo" in para.text
        assert "Project Name" not in para.text

        # Verify the output is a valid docx
        output = tmp_path / "output.docx"
        doc.save(str(output))
        reopened = Document(str(output))
        assert "DocForge Demo" in reopened.paragraphs[marker.paragraph_index].text

    def test_removes_red_color(self, simple_placeholder_path):
        analysis = parse_template(simple_placeholder_path)
        doc = Document(str(simple_placeholder_path))
        store = _make_data_store_with_json()

        marker = next(
            m for m in analysis.markers if m.marker_type == MarkerType.VARIABLE_PLACEHOLDER
        )
        mapping = MappingEntry(
            marker_id=marker.id,
            data_source="config.json",
            path="project.name",
        )

        renderer = PlaceholderRenderer()
        renderer.render(marker, store, doc, mapping)

        # Check color was removed from the replaced run
        run = doc.paragraphs[marker.paragraph_index].runs[marker.run_indices[0]]
        assert run.font.color.rgb is None

    def test_missing_data_source(self, simple_placeholder_path):
        analysis = parse_template(simple_placeholder_path)
        doc = Document(str(simple_placeholder_path))
        store = DataStore()  # Empty store

        marker = next(
            m for m in analysis.markers if m.marker_type == MarkerType.VARIABLE_PLACEHOLDER
        )
        mapping = MappingEntry(
            marker_id=marker.id,
            data_source="nonexistent.csv",
            field="Name",
        )

        renderer = PlaceholderRenderer()
        result = renderer.render(marker, store, doc, mapping)
        assert result.success is False
        assert "Could not resolve" in result.error


class TestTableRenderer:
    def test_can_handle(self):
        renderer = TableRenderer()
        sample = TemplateMarker(
            id="m-0",
            text="data",
            marker_type=MarkerType.SAMPLE_DATA,
            paragraph_index=-1,
            run_indices=[0],
            table_id="table-0",
            row_index=1,
        )
        placeholder = TemplateMarker(
            id="m-1",
            text="Name",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[0],
        )
        assert renderer.can_handle(sample) is True
        assert renderer.can_handle(placeholder) is False

    def test_render_table_direct(self, simple_table_path, tmp_path):
        doc = Document(str(simple_table_path))
        store = load_data_sources([FIXTURES_DIR / "data" / "metrics.xlsx"])

        mapping = MappingEntry(
            marker_id="table-0",
            data_source="metrics.xlsx",
            sheet="KPIs",
        )

        result = render_table_direct("table-0", store, doc, mapping)
        assert result.success is True

        # Table should now have header + 3 data rows
        table = doc.tables[0]
        assert len(table.rows) == 4  # 1 header + 3 KPI rows

        # Verify data in cells
        assert table.rows[1].cells[0].text == "Users"
        assert table.rows[1].cells[1].text == "5000"

        # Save and verify
        output = tmp_path / "output.docx"
        doc.save(str(output))

    def test_populates_with_csv_data(self, table_with_sample_data_path, tmp_path):
        analysis = parse_template(table_with_sample_data_path)
        doc = Document(str(table_with_sample_data_path))
        store = _make_data_store_with_csv()

        # Find a sample data marker
        sample_marker = next(m for m in analysis.markers if m.marker_type == MarkerType.SAMPLE_DATA)
        mapping = MappingEntry(
            marker_id=sample_marker.id,
            data_source="project_status.csv",
        )

        renderer = TableRenderer()
        result = renderer.render(sample_marker, store, doc, mapping)
        assert result.success is True

        # Table should now have header + 4 CSV rows (sample data removed, new data added)
        table = doc.tables[0]
        assert len(table.rows) == 5  # 1 header + 4 data rows
        assert table.rows[1].cells[0].text == "Alpha"

    def test_missing_data_source(self, table_with_sample_data_path):
        analysis = parse_template(table_with_sample_data_path)
        doc = Document(str(table_with_sample_data_path))
        store = DataStore()

        marker = next(m for m in analysis.markers if m.marker_type == MarkerType.SAMPLE_DATA)
        mapping = MappingEntry(
            marker_id=marker.id,
            data_source="nonexistent.csv",
        )

        renderer = TableRenderer()
        result = renderer.render(marker, store, doc, mapping)
        assert result.success is False
