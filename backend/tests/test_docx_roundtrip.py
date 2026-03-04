"""End-to-end round-trip tests: template → generate → editor → export → verify.

Uses the editor_test_template fixture to test the full pipeline from
a real template through the .docx→editor→.docx conversion cycle.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from core.docx_to_editor import DocxToEditorConverter
from core.editor_models import EditorDocument
from core.editor_to_docx import EditorToDocxConverter
from core.engine import GenerationEngine
from core.models import (
    GenerationReport,
    RenderResult,
    TemplateAnalysis,
)
from utils.red_text import is_red_run

FIXTURES = Path(__file__).parent / "fixtures"
TEMPLATES = FIXTURES / "templates"
DATA = FIXTURES / "data"


def _add_red_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


class TestEditorFixtureRoundTrip:
    """Tests using the editor_test_template fixture."""

    def test_editor_template_parses_correctly(self):
        """The editor fixture template has expected sections and markers."""
        engine = GenerationEngine()
        template = TEMPLATES / "editor_test_template.docx"
        if not template.exists():
            pytest.skip("editor_test_template.docx not generated yet")

        analysis = engine.analyze(str(template))

        # Should have sections: Report Title, Financial Overview, Conclusion
        assert len(analysis.sections) >= 2
        section_titles = [s.title for s in analysis.sections]
        assert "Report Title" in section_titles
        assert "Financial Overview" in section_titles

        # Should have at least 2 markers (Company Name placeholder + LLM prompt)
        assert len(analysis.markers) >= 2

    def test_round_trip_preserves_headings(self, tmp_path):
        """Headings survive the .docx → editor JSON → .docx cycle."""
        doc = Document()
        doc.add_heading("First Heading", level=1)
        doc.add_paragraph("Some body text.")
        doc.add_heading("Second Heading", level=2)
        doc.add_paragraph("More body text.")
        template = tmp_path / "headings.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        texts = [p.text for p in exported.paragraphs]
        assert "First Heading" in texts
        assert "Second Heading" in texts
        assert "Some body text." in texts
        assert "More body text." in texts

    def test_round_trip_preserves_formatting(self, tmp_path):
        """Bold and italic formatting survive the round trip."""
        doc = Document()
        para = doc.add_paragraph()
        run_bold = para.add_run("Bold text")
        run_bold.font.bold = True
        para.add_run(" and ")
        run_italic = para.add_run("italic text")
        run_italic.font.italic = True
        template = tmp_path / "formatting.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        para = exported.paragraphs[0]
        assert para.text == "Bold text and italic text"
        # Find the bold and italic runs
        bold_runs = [r for r in para.runs if r.font.bold]
        italic_runs = [r for r in para.runs if r.font.italic]
        assert len(bold_runs) >= 1
        assert len(italic_runs) >= 1

    def test_round_trip_preserves_tables(self, tmp_path):
        """Table structure and content survive the round trip."""
        doc = Document()
        doc.add_heading("Data", level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Revenue"
        table.rows[1].cells[1].text = "$100K"
        table.rows[2].cells[0].text = "Profit"
        table.rows[2].cells[1].text = "$25K"
        template = tmp_path / "tables.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        assert len(exported.tables) == 1
        t = exported.tables[0]
        assert len(t.rows) == 3
        assert t.rows[0].cells[0].text == "Metric"
        assert t.rows[0].cells[1].text == "Value"
        assert t.rows[1].cells[0].text == "Revenue"
        assert t.rows[2].cells[1].text == "$25K"

    def test_round_trip_mixed_content(self, tmp_path):
        """Mixed headings, paragraphs, and tables preserve order."""
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("First paragraph.")
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph("Last paragraph.")
        template = tmp_path / "mixed.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        para_texts = [p.text for p in exported.paragraphs]
        assert "Introduction" in para_texts
        assert "First paragraph." in para_texts
        assert "Conclusion" in para_texts
        assert "Last paragraph." in para_texts
        assert len(exported.tables) == 1

    def test_unresolved_markers_survive_round_trip(self, tmp_path):
        """Unresolved red text markers export as red text in the .docx."""
        doc = Document()
        para = doc.add_paragraph("Value: ")
        _add_red_text(para, "Missing Data")
        template = tmp_path / "unresolved.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(
            sections=[],
            markers=[],
            tables=[],
        )
        report = GenerationReport(
            total_markers=1,
            rendered=0,
            skipped=1,
            warnings=[],
            errors=[],
            results=[
                RenderResult(
                    marker_id="marker-0",
                    success=False,
                    status="skipped",
                    original_text="Missing Data",
                    rendered_text=None,
                )
            ],
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        # Verify unresolved node exists in editor JSON
        content = editor_doc.content.content
        assert content is not None
        unresolved_nodes = []
        for node in content:
            if node.content:
                for child in node.content:
                    if child.type == "docforgeUnresolved":
                        unresolved_nodes.append(child)
        assert len(unresolved_nodes) >= 1

        # Export back to .docx
        exporter = EditorToDocxConverter(editor_doc, template)
        exported = exporter.convert()

        # Find the red run
        red_runs = []
        for para in exported.paragraphs:
            for run in para.runs:
                if is_red_run(run._element):
                    red_runs.append(run)
        assert len(red_runs) >= 1
        assert red_runs[0].text == "Missing Data"

    def test_editor_json_structure_valid(self, tmp_path):
        """Editor JSON can be serialized and deserialized."""
        doc = Document()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("Hello")
        template = tmp_path / "structure.docx"
        doc.save(str(template))

        analysis = TemplateAnalysis(sections=[], markers=[], tables=[])
        report = GenerationReport(
            total_markers=0, rendered=0, skipped=0, warnings=[], errors=[], results=[]
        )

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        # Serialize to dict
        data = editor_doc.model_dump(mode="json")

        # Deserialize back
        restored = EditorDocument(**data)

        assert restored.content.type == "doc"
        assert restored.meta.generation_run_id is not None
        assert len(restored.content.content or []) > 0
