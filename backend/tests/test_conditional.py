"""Tests for conditional section evaluation and removal."""

from pathlib import Path

import pandas as pd
from docx import Document

from core.conditional import evaluate_condition, remove_section_content
from core.data_loader import DataStore
from core.models import (
    ConditionalConfig,
    MarkerType,
    Section,
    TemplateAnalysis,
    TemplateMarker,
)
from extractors.base import ExtractedData


def _make_store_with_data(name: str, columns: dict) -> DataStore:
    store = DataStore()
    df = pd.DataFrame(columns)
    store.add(name, ExtractedData(source_path=Path(name), dataframes={"default": df}))
    return store


def _make_store_with_text(name: str, text: str) -> DataStore:
    store = DataStore()
    store.add(
        name,
        ExtractedData(source_path=Path(name), text_content=text),
    )
    return store


class TestDataPresenceTrue:
    def test_data_presence_true(self):
        """Source exists with data -> True."""
        store = _make_store_with_data("report.csv", {"Sales": [100, 200]})
        config = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="report.csv",
            include=True,
        )

        result = evaluate_condition(config, store)
        assert result is True


class TestDataPresenceFalseMissing:
    def test_data_presence_false_missing(self):
        """Source doesn't exist -> False."""
        store = DataStore()
        config = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="nonexistent.csv",
            include=True,
        )

        result = evaluate_condition(config, store)
        assert result is False


class TestDataPresenceFieldCheck:
    def test_data_presence_field_check(self):
        """Check specific field presence: existing field -> True, missing field -> False."""
        store = _make_store_with_data("report.csv", {"Sales": [100], "Region": ["US"]})

        config_exists = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="report.csv",
            field="Sales",
            include=True,
        )
        assert evaluate_condition(config_exists, store) is True

        config_missing = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="report.csv",
            field="NonexistentField",
            include=True,
        )
        assert evaluate_condition(config_missing, store) is False


class TestExplicitEquals:
    def test_explicit_equals(self):
        """Explicit equals condition checks first row value."""
        store = _make_store_with_data("report.csv", {"Status": ["Active"]})

        config_match = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="report.csv",
            field="Status",
            operator="equals",
            value="Active",
            include=True,
        )
        assert evaluate_condition(config_match, store) is True

        config_no_match = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="report.csv",
            field="Status",
            operator="equals",
            value="Inactive",
            include=True,
        )
        assert evaluate_condition(config_no_match, store) is False


class TestExplicitGt:
    def test_explicit_gt(self):
        """Numeric greater-than condition."""
        store = _make_store_with_data("metrics.csv", {"Score": [85]})

        config_gt = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="metrics.csv",
            field="Score",
            operator="gt",
            value="70",
            include=True,
        )
        assert evaluate_condition(config_gt, store) is True

        config_gt_fail = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="metrics.csv",
            field="Score",
            operator="gt",
            value="90",
            include=True,
        )
        assert evaluate_condition(config_gt_fail, store) is False


class TestExplicitContains:
    def test_explicit_contains(self):
        """String contains condition."""
        store = _make_store_with_data(
            "report.csv", {"Description": ["Quarterly revenue report for Q3"]}
        )

        config_contains = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="report.csv",
            field="Description",
            operator="contains",
            value="revenue",
            include=True,
        )
        assert evaluate_condition(config_contains, store) is True

        config_no_contains = ConditionalConfig(
            section_id="s1",
            condition_type="explicit",
            data_source="report.csv",
            field="Description",
            operator="contains",
            value="annual",
            include=True,
        )
        assert evaluate_condition(config_no_contains, store) is False


class TestIncludeFalseInverts:
    def test_include_false_inverts(self):
        """include=False inverts the result: present source -> False (excluded)."""
        store = _make_store_with_data("report.csv", {"Sales": [100]})

        # With include=True, data_presence of existing source -> True
        config_include = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="report.csv",
            include=True,
        )
        assert evaluate_condition(config_include, store) is True

        # With include=False, data_presence of existing source -> False (inverted)
        config_exclude = ConditionalConfig(
            section_id="s1",
            condition_type="data_presence",
            data_source="report.csv",
            include=False,
        )
        assert evaluate_condition(config_exclude, store) is False


class TestRemoveSection:
    def test_remove_section(self):
        """Create a 3-section document, remove middle section, verify paragraphs removed."""
        doc = Document()

        # Section 1: paragraph index 0
        doc.add_paragraph("Section 1 content")

        # Section 2: paragraph indices 1, 2
        doc.add_paragraph("Section 2 heading")
        doc.add_paragraph("Section 2 body text")

        # Section 3: paragraph index 3
        doc.add_paragraph("Section 3 content")

        marker_in_s2 = TemplateMarker(
            id="m_s2",
            text="Some marker",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=1,
            run_indices=[0],
            section_id="s2",
        )

        sections = [
            Section(id="s1", title="Sec 1", level=1, paragraph_index=0, markers=[]),
            Section(
                id="s2",
                title="Sec 2",
                level=1,
                paragraph_index=1,
                markers=[marker_in_s2],
            ),
            Section(id="s3", title="Sec 3", level=1, paragraph_index=3, markers=[]),
        ]
        analysis = TemplateAnalysis(sections=sections, markers=[marker_in_s2], tables=[])

        # Before removal: 4 paragraphs
        assert len(doc.paragraphs) == 4

        removed_marker_ids = remove_section_content(doc, "s2", analysis)

        # Middle section (2 paragraphs) removed -> 2 remaining
        assert len(doc.paragraphs) == 2
        remaining_texts = [p.text for p in doc.paragraphs]
        assert "Section 1 content" in remaining_texts
        assert "Section 3 content" in remaining_texts
        assert "Section 2 heading" not in remaining_texts
        assert "Section 2 body text" not in remaining_texts

        # Removed markers include the marker from section 2
        assert "m_s2" in removed_marker_ids
