"""Tests for the .docx to TipTap JSON converter."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor

from core.docx_to_editor import DocxToEditorConverter
from core.models import (
    AutoResolutionMatch,
    AutoResolutionReport,
    GenerationReport,
    MappingEntry,
    MarkerType,
    RenderResult,
    Section,
    TemplateAnalysis,
    TemplateMarker,
)


def _add_red_text(paragraph, text):
    """Add a red-colored run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return run


def _make_analysis(markers, sections=None, tables=None):
    return TemplateAnalysis(
        sections=sections or [],
        markers=markers,
        tables=tables or [],
    )


def _make_report(results, total=None):
    rendered = sum(1 for r in results if r.success)
    return GenerationReport(
        total_markers=total or len(results),
        rendered=rendered,
        skipped=len(results) - rendered,
        warnings=[],
        errors=[],
        results=results,
    )


class TestBasicConversion:
    def test_empty_document(self):
        doc = Document()
        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()
        assert result.content.type == "doc"
        assert result.content.content == []

    def test_single_paragraph(self):
        doc = Document()
        doc.add_paragraph("Hello world")

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        assert result.content.type == "doc"
        assert len(result.content.content) == 1
        para_node = result.content.content[0]
        assert para_node.type == "paragraph"
        assert len(para_node.content) == 1
        assert para_node.content[0].type == "text"
        assert para_node.content[0].text == "Hello world"

    def test_heading_conversion(self):
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Subtitle", level=2)

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        assert len(result.content.content) == 2
        h1 = result.content.content[0]
        assert h1.type == "heading"
        assert h1.attrs["level"] == 1
        h2 = result.content.content[1]
        assert h2.type == "heading"
        assert h2.attrs["level"] == 2


class TestFormattingMarks:
    def test_bold_text(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.font.bold = True

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        assert text_node.text == "Bold text"
        assert any(m.type == "bold" for m in text_node.marks)

    def test_italic_text(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.font.italic = True

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        assert any(m.type == "italic" for m in text_node.marks)

    def test_underline_text(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Underlined")
        run.font.underline = True

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        assert any(m.type == "underline" for m in text_node.marks)

    def test_font_size_and_name(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Styled text")
        run.font.size = Pt(14)
        run.font.name = "Arial"

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        style_marks = [m for m in text_node.marks if m.type == "textStyle"]
        assert len(style_marks) == 1
        assert style_marks[0].attrs["fontSize"] == "14.0pt"
        assert style_marks[0].attrs["fontFamily"] == "Arial"

    def test_font_color(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Colored")
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        style_marks = [m for m in text_node.marks if m.type == "textStyle"]
        assert len(style_marks) == 1
        assert style_marks[0].attrs["color"] == "#008000"

    def test_multiple_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold and italic")
        run.font.bold = True
        run.font.italic = True

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        text_node = result.content.content[0].content[0]
        mark_types = {m.type for m in text_node.marks}
        assert "bold" in mark_types
        assert "italic" in mark_types


class TestRenderedContent:
    def test_placeholder_rendered_high_confidence(self):
        """Placeholders with high confidence should NOT get docforgeRendered mark."""
        doc = Document()
        doc.add_heading("Intro", level=1)
        para = doc.add_paragraph("The project is ")
        para.add_run("DocForge")  # Rendered placeholder (red removed)
        para.add_run(".")

        marker = TemplateMarker(
            id="marker-0",
            text="Project Name",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            section_id="section-0",
            paragraph_index=1,
            run_indices=[0],
        )
        section = Section(
            id="section-0", title="Intro", level=1, paragraph_index=0, markers=[marker]
        )
        analysis = _make_analysis([marker], sections=[section])
        result_entry = RenderResult(marker_id="marker-0", success=True, rendered_by="placeholder")
        report = _make_report([result_entry])
        auto_report = AutoResolutionReport(
            matches=[
                AutoResolutionMatch(
                    marker_id="marker-0",
                    data_source="config.json",
                    field="name",
                    confidence=0.95,
                    match_type="exact",
                    reasoning="Exact match",
                )
            ],
            unresolved=[],
        )

        converter = DocxToEditorConverter(
            doc, report, analysis, [], auto_resolution_report=auto_report
        )
        editor_doc = converter.convert()

        # The rendered text should NOT have docforgeRendered mark (confidence >= 0.8)
        para_node = editor_doc.content.content[1]
        first_text = para_node.content[0]
        assert first_text.text == "The project is "
        second_text = para_node.content[1]
        assert second_text.text == "DocForge"
        has_rendered_mark = second_text.marks and any(
            m.type == "docforgeRendered" for m in second_text.marks
        )
        assert not has_rendered_mark

    def test_placeholder_rendered_low_confidence(self):
        """Placeholders with low confidence should get docforgeRendered mark."""
        doc = Document()
        para = doc.add_paragraph("Value: ")
        para.add_run("42")  # Rendered but low confidence

        marker = TemplateMarker(
            id="marker-0",
            text="Amount",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[1],
        )
        analysis = _make_analysis([marker])
        result_entry = RenderResult(marker_id="marker-0", success=True, rendered_by="placeholder")
        report = _make_report([result_entry])
        auto_report = AutoResolutionReport(
            matches=[
                AutoResolutionMatch(
                    marker_id="marker-0",
                    data_source="data.xlsx",
                    field="amount",
                    confidence=0.6,
                    match_type="fuzzy",
                    reasoning="Fuzzy match",
                )
            ],
            unresolved=[],
        )

        converter = DocxToEditorConverter(
            doc, report, analysis, [], auto_resolution_report=auto_report
        )
        editor_doc = converter.convert()

        para_node = editor_doc.content.content[0]
        rendered_text = para_node.content[1]
        assert rendered_text.text == "42"
        rendered_marks = [m for m in (rendered_text.marks or []) if m.type == "docforgeRendered"]
        assert len(rendered_marks) == 1
        assert rendered_marks[0].attrs["markerId"] == "marker-0"
        assert rendered_marks[0].attrs["confidence"] == 0.6

    def test_llm_rendered_content(self):
        """LLM-rendered content should always get docforgeRendered mark."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("LLM generated executive summary text")

        marker = TemplateMarker(
            id="marker-0",
            text="Provide an executive summary",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[0],
        )
        analysis = _make_analysis([marker])
        result_entry = RenderResult(
            marker_id="marker-0",
            success=True,
            rendered_by="llm",
            llm_usage={"model": "gpt-4", "prompt_tokens": 100, "completion_tokens": 50},
        )
        report = _make_report([result_entry])

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        para_node = editor_doc.content.content[0]
        text_node = para_node.content[0]
        rendered_marks = [m for m in text_node.marks if m.type == "docforgeRendered"]
        assert len(rendered_marks) == 1
        assert rendered_marks[0].attrs["renderType"] == "llm"

    def test_unresolved_marker(self):
        """Unresolved red text markers become docforgeUnresolved nodes."""
        doc = Document()
        para = doc.add_paragraph("Start ")
        _add_red_text(para, "Unresolved Prompt")
        para.add_run(" end")

        marker = TemplateMarker(
            id="marker-0",
            text="Unresolved Prompt",
            marker_type=MarkerType.LLM_PROMPT,
            paragraph_index=0,
            run_indices=[1],
        )
        analysis = _make_analysis([marker])
        result_entry = RenderResult(marker_id="marker-0", success=False, error="No mapping")
        report = _make_report([result_entry])

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        para_node = editor_doc.content.content[0]
        assert len(para_node.content) == 3
        assert para_node.content[0].text == "Start "
        assert para_node.content[1].type == "docforgeUnresolved"
        assert para_node.content[1].attrs["markerId"] == "marker-0"
        assert para_node.content[1].attrs["originalText"] == "Unresolved Prompt"
        assert para_node.content[2].text == " end"


class TestTableConversion:
    def test_basic_table(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        for i, header in enumerate(["Name", "Value", "Unit"]):
            table.rows[0].cells[i].text = header
        table.rows[1].cells[0].text = "Revenue"
        table.rows[1].cells[1].text = "100000"
        table.rows[1].cells[2].text = "USD"

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        table_node = result.content.content[0]
        assert table_node.type == "table"
        assert table_node.attrs["docforgeTableId"] == "table-0"
        assert len(table_node.content) == 2  # 2 rows

        # Header row
        header_row = table_node.content[0]
        assert header_row.type == "tableRow"
        assert len(header_row.content) == 3
        assert header_row.content[0].type == "tableHeader"

        # Data row
        data_row = table_node.content[1]
        assert data_row.content[0].type == "tableCell"


class TestDocumentInterleaving:
    def test_paragraph_table_paragraph(self):
        """Paragraphs and tables appear in correct document order."""
        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        doc.add_paragraph("After table")

        analysis = _make_analysis([])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        types = [n.type for n in result.content.content]
        assert types == ["paragraph", "table", "paragraph"]

    def test_heading_with_section_id(self):
        """Heading nodes carry section IDs for navigation."""
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Some text")

        section = Section(id="section-0", title="Introduction", level=1, paragraph_index=0)
        analysis = _make_analysis([], sections=[section])
        report = _make_report([])
        converter = DocxToEditorConverter(doc, report, analysis, [])
        result = converter.convert()

        heading = result.content.content[0]
        assert heading.type == "heading"
        assert heading.attrs["id"] == "section-section-0"


class TestMarkerMetadata:
    def test_metadata_built_for_all_markers(self):
        doc = Document()
        doc.add_paragraph("Text")

        markers = [
            TemplateMarker(
                id="m-0",
                text="Author",
                marker_type=MarkerType.VARIABLE_PLACEHOLDER,
                paragraph_index=0,
                run_indices=[0],
            ),
            TemplateMarker(
                id="m-1",
                text="Summarize the report",
                marker_type=MarkerType.LLM_PROMPT,
                paragraph_index=0,
                run_indices=[0],
            ),
        ]
        analysis = _make_analysis(markers)
        results = [
            RenderResult(marker_id="m-0", success=True, rendered_by="placeholder"),
            RenderResult(
                marker_id="m-1",
                success=True,
                rendered_by="llm",
                llm_usage={"model": "gpt-4", "prompt_tokens": 200, "completion_tokens": 80},
            ),
        ]
        report = _make_report(results)
        mapping = MappingEntry(marker_id="m-0", data_source="config.json", field="author")

        converter = DocxToEditorConverter(doc, report, analysis, [mapping])
        editor_doc = converter.convert()

        meta = editor_doc.meta.marker_metadata
        assert "m-0" in meta
        assert meta["m-0"].rendered_by == "placeholder"
        assert meta["m-0"].original_text == "Author"
        assert meta["m-0"].mapping_snapshot is not None

        assert "m-1" in meta
        assert meta["m-1"].rendered_by == "llm"
        assert meta["m-1"].llm_model == "gpt-4"
        assert meta["m-1"].llm_prompt_tokens == 200

    def test_unresolved_marker_metadata(self):
        doc = Document()
        doc.add_paragraph("Text")

        marker = TemplateMarker(
            id="m-0",
            text="Missing Data",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[0],
        )
        analysis = _make_analysis([marker])
        report = _make_report([RenderResult(marker_id="m-0", success=False, error="No mapping")])

        converter = DocxToEditorConverter(doc, report, analysis, [])
        editor_doc = converter.convert()

        meta = editor_doc.meta.marker_metadata
        assert meta["m-0"].rendered_by == "unresolved"


class TestEditorDocumentMeta:
    def test_meta_fields(self):
        doc = Document()
        analysis = _make_analysis([])
        report = _make_report([])

        converter = DocxToEditorConverter(doc, report, analysis, [], project_id=42, run_id=7)
        editor_doc = converter.convert()

        assert editor_doc.meta.project_id == 42
        assert editor_doc.meta.generation_run_id == 7
        assert editor_doc.meta.template_analysis is not None
