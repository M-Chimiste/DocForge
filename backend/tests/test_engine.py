"""Tests for GenerationEngine — full pipeline integration."""

import pytest
from docx import Document

from core.engine import GenerationEngine
from core.models import MappingEntry


@pytest.fixture
def engine():
    return GenerationEngine()


class TestAnalyze:
    def test_analyze_simple_placeholder(self, engine, templates_dir):
        analysis = engine.analyze(templates_dir / "simple_placeholder.docx")
        assert len(analysis.sections) >= 1
        assert len(analysis.markers) >= 1
        assert analysis.markers[0].text == "Project Name"

    def test_analyze_mixed_markers(self, engine, templates_dir):
        analysis = engine.analyze(templates_dir / "mixed_markers.docx")
        assert len(analysis.sections) >= 3
        assert len(analysis.markers) >= 3
        assert len(analysis.tables) >= 1


class TestGenerate:
    def test_placeholder_substitution(self, engine, templates_dir, data_dir, tmp_path):
        """End-to-end: replace a placeholder with a JSON value."""
        output = tmp_path / "output.docx"
        mappings = [
            MappingEntry(
                marker_id="marker-0",
                data_source="config.json",
                field="name",
                path="project",
            ),
        ]
        report = engine.generate(
            templates_dir / "simple_placeholder.docx",
            [data_dir / "config.json"],
            mappings,
            output,
        )
        assert output.exists()
        # Check at least one succeeded
        succeeded = [r for r in report.results if r.success]
        assert len(succeeded) >= 1

        # Verify the text was actually substituted in the output document
        doc = Document(str(output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "DocForge Demo" in full_text

    def test_table_population(self, engine, templates_dir, data_dir, tmp_path):
        """End-to-end: populate a skeleton table from CSV data."""
        output = tmp_path / "output.docx"
        mappings = [
            MappingEntry(
                marker_id="table-0",
                data_source="project_status.csv",
            ),
        ]
        engine.generate(
            templates_dir / "simple_table.docx",
            [data_dir / "project_status.csv"],
            mappings,
            output,
        )
        assert output.exists()

        # Verify the table has data rows
        doc = Document(str(output))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        # Header row + at least some data rows
        assert len(table.rows) >= 2

    def test_missing_mapping_fail_forward(self, engine, templates_dir, data_dir, tmp_path):
        """Markers without mappings should fail gracefully, not crash."""
        output = tmp_path / "output.docx"
        # Provide no mappings at all
        report = engine.generate(
            templates_dir / "simple_placeholder.docx",
            [data_dir / "config.json"],
            [],
            output,
        )
        assert output.exists()
        # All markers should report failure
        failed = [r for r in report.results if not r.success]
        assert len(failed) >= 1
        assert any("No mapping" in r.error for r in failed)

    def test_mixed_template_generation(self, engine, templates_dir, data_dir, tmp_path):
        """Mixed template with placeholders, tables, and LLM prompts."""
        output = tmp_path / "output.docx"
        mappings = [
            MappingEntry(
                marker_id="marker-0",
                data_source="config.json",
                field="author",
                path="settings",
            ),
            MappingEntry(
                marker_id="marker-1",
                data_source="config.json",
                field="date",
                path="project",
            ),
            MappingEntry(
                marker_id="table-0",
                data_source="metrics.xlsx",
                sheet="Revenue",
            ),
        ]
        report = engine.generate(
            templates_dir / "mixed_markers.docx",
            [data_dir / "config.json", data_dir / "metrics.xlsx"],
            mappings,
            output,
        )
        assert output.exists()
        # Some should succeed, LLM prompt should fail (no mapping)
        succeeded = [r for r in report.results if r.success]
        assert len(succeeded) >= 1

    def test_output_dir_created(self, engine, templates_dir, data_dir, tmp_path):
        """Output directory should be created if it doesn't exist."""
        output = tmp_path / "nested" / "dir" / "output.docx"
        engine.generate(
            templates_dir / "simple_placeholder.docx",
            [data_dir / "config.json"],
            [],
            output,
        )
        assert output.exists()
