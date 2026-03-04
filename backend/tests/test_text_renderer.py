"""Tests for the TextRenderer: inject text content into LLM_PROMPT markers."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import RGBColor

from core.data_loader import DataStore
from core.models import MappingEntry, MarkerType, TemplateMarker
from extractors.base import ExtractedData
from renderers.text_renderer import TextRenderer


def _make_doc_with_red_run(text: str) -> Document:
    """Create a document with a single paragraph containing red text."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return doc


def _make_llm_marker(
    marker_id: str = "m1",
    text: str = "Analyze this section",
    paragraph_index: int = 0,
    run_indices: list[int] | None = None,
) -> TemplateMarker:
    return TemplateMarker(
        id=marker_id,
        text=text,
        marker_type=MarkerType.LLM_PROMPT,
        paragraph_index=paragraph_index,
        run_indices=run_indices or [0],
    )


def _make_variable_marker(marker_id: str = "m1") -> TemplateMarker:
    return TemplateMarker(
        id=marker_id,
        text="Author",
        marker_type=MarkerType.VARIABLE_PLACEHOLDER,
        paragraph_index=0,
        run_indices=[0],
    )


def _make_store_with_text(name: str, text: str) -> DataStore:
    store = DataStore()
    store.add(
        name,
        ExtractedData(source_path=Path(name), text_content=text),
    )
    return store


class TestCanHandleLlmPrompt:
    def test_can_handle_llm_prompt(self):
        """TextRenderer handles LLM_PROMPT markers, not VARIABLE_PLACEHOLDER."""
        renderer = TextRenderer()

        llm_marker = _make_llm_marker()
        assert renderer.can_handle(llm_marker) is True

        var_marker = _make_variable_marker()
        assert renderer.can_handle(var_marker) is False


class TestInjectSingleParagraph:
    def test_inject_single_paragraph(self):
        """Create docx with red LLM_PROMPT text, render with text content, verify replacement."""
        doc = _make_doc_with_red_run("Summarize the project")
        marker = _make_llm_marker(paragraph_index=0, run_indices=[0])
        store = _make_store_with_text("summary.txt", "This is the project summary.")
        mapping = MappingEntry(marker_id="m1", data_source="summary.txt")

        renderer = TextRenderer()
        result = renderer.render(marker, store, doc, mapping)

        assert result.success is True
        assert result.marker_id == "m1"

        # Text should be replaced
        para = doc.paragraphs[0]
        assert para.runs[0].text == "This is the project summary."

        # Red color should be cleared (rgb set to None)
        assert para.runs[0].font.color.rgb is None


class TestInjectMultiParagraph:
    def test_inject_multi_paragraph(self):
        """Text with newlines creates additional paragraphs."""
        doc = _make_doc_with_red_run("Summarize the project")
        marker = _make_llm_marker(paragraph_index=0, run_indices=[0])
        multi_text = "First paragraph.\nSecond paragraph.\nThird paragraph."
        store = _make_store_with_text("content.txt", multi_text)
        mapping = MappingEntry(marker_id="m1", data_source="content.txt")

        renderer = TextRenderer()
        result = renderer.render(marker, store, doc, mapping)

        assert result.success is True

        # The first paragraph should contain the first line
        assert doc.paragraphs[0].runs[0].text == "First paragraph."

        # Additional paragraphs should have been inserted
        all_texts = [p.text for p in doc.paragraphs]
        assert "Second paragraph." in all_texts
        assert "Third paragraph." in all_texts


class TestMissingTextSource:
    def test_missing_text_source(self):
        """No text content available -> failure result."""
        doc = _make_doc_with_red_run("Summarize the project")
        marker = _make_llm_marker(paragraph_index=0, run_indices=[0])

        # Store with a source that has no text content and no 'content' column
        store = DataStore()
        df = pd.DataFrame({"col1": [1]})
        store.add(
            "numbers.csv",
            ExtractedData(source_path=Path("numbers.csv"), dataframes={"default": df}),
        )
        mapping = MappingEntry(marker_id="m1", data_source="numbers.csv")

        renderer = TextRenderer()
        result = renderer.render(marker, store, doc, mapping)

        assert result.success is False
        assert result.marker_id == "m1"
        assert "no text content" in result.error.lower()
