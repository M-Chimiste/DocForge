"""Tests for mapping and output validation functions."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import RGBColor

from core.data_loader import DataStore
from core.models import (
    MappingEntry,
    MarkerType,
    TemplateAnalysis,
    TemplateMarker,
)
from core.validators import validate_mappings, validate_output
from extractors.base import ExtractedData


def _make_marker(
    marker_id: str,
    text: str,
    marker_type: MarkerType = MarkerType.VARIABLE_PLACEHOLDER,
    table_id: str | None = None,
) -> TemplateMarker:
    return TemplateMarker(
        id=marker_id,
        text=text,
        marker_type=marker_type,
        paragraph_index=0,
        run_indices=[0],
        table_id=table_id,
    )


def _make_analysis(markers: list[TemplateMarker]) -> TemplateAnalysis:
    return TemplateAnalysis(sections=[], markers=markers, tables=[])


def _make_store_with_csv(name: str, columns: dict) -> DataStore:
    store = DataStore()
    df = pd.DataFrame(columns)
    store.add(name, ExtractedData(source_path=Path(name), dataframes={"default": df}))
    return store


class TestAllMappedNoIssues:
    def test_all_mapped_no_issues(self):
        """All markers have valid mappings -> empty issues list."""
        store = _make_store_with_csv("data.csv", {"Author": ["Alice"], "Title": ["Doc"]})
        markers = [_make_marker("m1", "Author"), _make_marker("m2", "Title")]
        analysis = _make_analysis(markers)

        mappings = [
            MappingEntry(marker_id="m1", data_source="data.csv", field="Author"),
            MappingEntry(marker_id="m2", data_source="data.csv", field="Title"),
        ]

        issues = validate_mappings(analysis, mappings, store)
        assert issues == []


class TestUnmappedMarkerWarning:
    def test_unmapped_marker_warning(self):
        """One marker with no mapping -> warning."""
        store = _make_store_with_csv("data.csv", {"Author": ["Alice"]})
        markers = [_make_marker("m1", "Author"), _make_marker("m2", "Title")]
        analysis = _make_analysis(markers)

        mappings = [
            MappingEntry(marker_id="m1", data_source="data.csv", field="Author"),
        ]

        issues = validate_mappings(analysis, mappings, store)
        assert len(issues) == 1
        assert issues[0].level == "warning"
        assert issues[0].marker_id == "m2"
        assert "no mapping" in issues[0].message.lower()


class TestMissingSourceError:
    def test_missing_source_error(self):
        """Mapping references nonexistent source -> error."""
        store = _make_store_with_csv("data.csv", {"Author": ["Alice"]})
        markers = [_make_marker("m1", "Author")]
        analysis = _make_analysis(markers)

        mappings = [
            MappingEntry(marker_id="m1", data_source="nonexistent.csv", field="Author"),
        ]

        issues = validate_mappings(analysis, mappings, store)
        assert len(issues) == 1
        assert issues[0].level == "error"
        assert issues[0].marker_id == "m1"
        assert "not found" in issues[0].message.lower()


class TestMissingFieldError:
    def test_missing_field_error(self):
        """Mapping references nonexistent field -> error."""
        store = _make_store_with_csv("data.csv", {"Author": ["Alice"]})
        markers = [_make_marker("m1", "Author")]
        analysis = _make_analysis(markers)

        mappings = [
            MappingEntry(marker_id="m1", data_source="data.csv", field="NonexistentField"),
        ]

        issues = validate_mappings(analysis, mappings, store)
        assert len(issues) == 1
        assert issues[0].level == "error"
        assert issues[0].marker_id == "m1"
        assert "NonexistentField" in issues[0].message


class TestValidateOutputClean:
    def test_validate_output_clean(self):
        """Document with no red text -> no issues."""
        doc = Document()
        doc.add_paragraph("Normal text here")
        doc.add_paragraph("Another paragraph")

        issues = validate_output(doc)
        assert issues == []


class TestValidateOutputUnresolved:
    def test_validate_output_unresolved(self):
        """Document with red text -> warning about unresolved red text."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Unresolved placeholder")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        issues = validate_output(doc)
        assert len(issues) >= 1
        assert issues[0].level == "warning"
        assert "unresolved red text" in issues[0].message.lower()
