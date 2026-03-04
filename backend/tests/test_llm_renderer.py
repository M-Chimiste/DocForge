"""Tests for renderers.llm_renderer — LLMRenderer."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

from docx import Document
from docx.shared import RGBColor

from core.data_loader import DataStore
from core.llm_client import LLMClient, LLMConfig, LLMResponse
from core.models import (
    MappingEntry,
    MarkerType,
    Section,
    TemplateAnalysis,
    TemplateMarker,
)
from extractors.base import ExtractedData
from renderers.llm_renderer import LLMRenderer


def _make_doc_with_red_text(red_text: str, prefix: str = "Before ") -> Document:
    """Create a minimal Document with one paragraph containing red text."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(prefix)
    red_run = para.add_run(red_text)
    red_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return doc


def _make_marker(
    marker_id: str = "marker-0",
    text: str = "Summarize the findings",
    section_id: str | None = "section-0",
    paragraph_index: int = 0,
    run_indices: list[int] | None = None,
) -> TemplateMarker:
    return TemplateMarker(
        id=marker_id,
        text=text,
        marker_type=MarkerType.LLM_PROMPT,
        section_id=section_id,
        paragraph_index=paragraph_index,
        run_indices=run_indices or [1],
    )


def _make_analysis(markers=None, sections=None) -> TemplateAnalysis:
    return TemplateAnalysis(
        sections=sections or [],
        markers=markers or [],
        tables=[],
    )


def _make_data_store(**text_sources: str) -> DataStore:
    store = DataStore()
    for name, text in text_sources.items():
        store.add(
            name,
            ExtractedData(
                source_path=Path(name),
                dataframes={},
                metadata={},
                text_content=text,
            ),
        )
    return store


def _make_mock_client(response_content: str = "Generated LLM text") -> LLMClient:
    config = LLMConfig(provider="openai", model="gpt-4o-mini")
    client = LLMClient(config)
    client.complete = MagicMock(
        return_value=LLMResponse(
            content=response_content,
            model="gpt-4o-mini",
            prompt_tokens=50,
            completion_tokens=20,
            total_tokens=70,
        )
    )
    return client


# ---------------------------------------------------------------------------
# can_handle
# ---------------------------------------------------------------------------


class TestCanHandle:
    def test_handles_llm_prompt_when_configured(self):
        renderer = LLMRenderer()
        client = _make_mock_client()
        renderer.configure(client, None, None)
        marker = _make_marker()
        assert renderer.can_handle(marker) is True

    def test_rejects_when_no_client(self):
        renderer = LLMRenderer()
        marker = _make_marker()
        assert renderer.can_handle(marker) is False

    def test_rejects_when_unconfigured_client(self):
        renderer = LLMRenderer()
        config = LLMConfig()  # empty, not configured
        client = LLMClient(config)
        renderer.configure(client, None, None)
        marker = _make_marker()
        assert renderer.can_handle(marker) is False

    def test_rejects_non_llm_prompt(self):
        renderer = LLMRenderer()
        client = _make_mock_client()
        renderer.configure(client, None, None)
        marker = TemplateMarker(
            id="marker-0",
            text="Project Name",
            marker_type=MarkerType.VARIABLE_PLACEHOLDER,
            paragraph_index=0,
            run_indices=[0],
        )
        assert renderer.can_handle(marker) is False


# ---------------------------------------------------------------------------
# render
# ---------------------------------------------------------------------------


class TestRender:
    def test_basic_render(self):
        doc = _make_doc_with_red_text("Summarize the findings")
        marker = _make_marker(paragraph_index=0, run_indices=[1])
        section = Section(
            id="section-0", title="Findings", level=1, paragraph_index=0, markers=[marker]
        )
        analysis = _make_analysis(markers=[marker], sections=[section])
        store = _make_data_store(**{"report.txt": "The project succeeded."})
        mapping = MappingEntry(marker_id="marker-0", data_source="report.txt")

        client = _make_mock_client("The project achieved all its goals.")
        renderer = LLMRenderer()
        renderer.configure(client, analysis, [mapping])

        result = renderer.render(marker, store, doc, mapping)

        assert result.success is True
        assert result.llm_usage is not None
        assert result.llm_usage["prompt_tokens"] == 50
        assert result.llm_usage["model"] == "gpt-4o-mini"
        # Verify the LLM was called
        client.complete.assert_called_once()
        # Verify the prompt includes the instruction
        call_args = client.complete.call_args
        assert "Summarize the findings" in call_args[0][0]

    def test_context_includes_mapped_data(self):
        doc = _make_doc_with_red_text("Summarize the data")
        marker = _make_marker(paragraph_index=0, run_indices=[1])
        section = Section(
            id="section-0", title="Analysis", level=1, paragraph_index=0, markers=[marker]
        )
        analysis = _make_analysis(markers=[marker], sections=[section])
        store = _make_data_store(**{"notes.txt": "Important findings here."})
        mapping = MappingEntry(marker_id="marker-0", data_source="notes.txt")

        client = _make_mock_client("Summary of findings.")
        renderer = LLMRenderer()
        renderer.configure(client, analysis, [mapping])

        renderer.render(marker, store, doc, mapping)

        # Verify context data was included in the prompt
        prompt = client.complete.call_args[0][0]
        assert "notes.txt" in prompt
        assert "Important findings here" in prompt

    def test_fail_forward_on_llm_error(self):
        doc = _make_doc_with_red_text("Generate text")
        marker = _make_marker(paragraph_index=0, run_indices=[1])
        analysis = _make_analysis(markers=[marker])
        store = _make_data_store()
        mapping = MappingEntry(marker_id="marker-0", data_source="data.txt")

        config = LLMConfig(provider="openai", model="gpt-4o-mini")
        client = LLMClient(config)
        client.complete = MagicMock(side_effect=Exception("Rate limit exceeded"))

        renderer = LLMRenderer()
        renderer.configure(client, analysis, [mapping])

        result = renderer.render(marker, store, doc, mapping)

        assert result.success is False
        assert "Rate limit exceeded" in result.error

    def test_paragraph_out_of_range(self):
        doc = _make_doc_with_red_text("text")
        marker = _make_marker(paragraph_index=999, run_indices=[0])
        analysis = _make_analysis(markers=[marker])
        store = _make_data_store()
        mapping = MappingEntry(marker_id="marker-0", data_source="data.txt")

        client = _make_mock_client("output")
        renderer = LLMRenderer()
        renderer.configure(client, analysis, [mapping])

        result = renderer.render(marker, store, doc, mapping)

        assert result.success is False
        assert "out of range" in result.error

    def test_progress_callback_called(self):
        doc = _make_doc_with_red_text("Summarize")
        marker = _make_marker(paragraph_index=0, run_indices=[1])
        analysis = _make_analysis(markers=[marker])
        store = _make_data_store(**{"data.txt": "content"})
        mapping = MappingEntry(marker_id="marker-0", data_source="data.txt")

        client = _make_mock_client("output")
        callback = MagicMock()
        renderer = LLMRenderer()
        renderer.configure(client, analysis, [mapping], progress_callback=callback)

        renderer.render(marker, store, doc, mapping)

        assert callback.call_count >= 2
        # First call: llm_call_started
        first_call = callback.call_args_list[0][0][0]
        assert first_call["status"] == "llm_call_started"
        # Second call: llm_call_completed
        second_call = callback.call_args_list[1][0][0]
        assert second_call["status"] == "llm_call_completed"


# ---------------------------------------------------------------------------
# Engine integration: fallback to TextRenderer when no LLM
# ---------------------------------------------------------------------------


class TestFallbackToTextRenderer:
    def test_no_llm_config_uses_text_renderer(self, tmp_path, templates_dir, data_dir):
        """When no LLM is configured, LLM_PROMPT markers fall through to TextRenderer."""
        from core.engine import GenerationEngine

        engine = GenerationEngine()  # No llm_config
        # Engine should work exactly as before
        assert engine._llm_renderer.can_handle(_make_marker()) is False

    def test_with_llm_config_uses_llm_renderer(self):
        from core.engine import GenerationEngine

        config = LLMConfig(provider="openai", model="gpt-4o-mini")
        engine = GenerationEngine(llm_config=config)
        # Before generate(), renderer is not yet configured with a client
        # It gets configured during generate() call
        assert engine._llm_config == config
