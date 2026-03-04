"""Convert TipTap JSON editor state back to a .docx Document.

Uses the original template as a base to preserve headers, footers,
page layout, styles, and other formatting that TipTap cannot represent.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.editor_models import EditorDocument, EditorMark, EditorNode


class EditorToDocxConverter:
    """Convert TipTap JSON editor state back to a .docx Document."""

    def __init__(
        self,
        editor_document: EditorDocument,
        template_path: Path,
    ):
        self._editor = editor_document
        self._template_path = template_path

    def convert(self) -> Document:
        """Produce a python-docx Document from the editor state."""
        doc = Document(str(self._template_path))
        self._clear_body(doc)

        for node in self._editor.content.content or []:
            self._render_node(doc, node)

        return doc

    def _clear_body(self, doc: Document) -> None:
        """Remove all body content while preserving section properties."""
        body = doc.element.body
        children_to_remove = [child for child in body if child.tag != qn("w:sectPr")]
        for child in children_to_remove:
            body.remove(child)

    def _render_node(self, doc: Document, node: EditorNode) -> None:
        """Render a single editor node into the document."""
        if node.type == "heading":
            self._render_heading(doc, node)
        elif node.type == "paragraph":
            self._render_paragraph(doc, node)
        elif node.type == "table":
            self._render_table(doc, node)
        elif node.type == "docforgeUnresolved":
            self._render_unresolved_inline(doc, node)

    def _render_heading(self, doc: Document, node: EditorNode) -> None:
        level = node.attrs.get("level", 1) if node.attrs else 1
        paragraph = doc.add_heading("", level=level)
        paragraph.clear()
        self._render_inline_content(paragraph, node.content or [])

    def _render_paragraph(self, doc: Document, node: EditorNode) -> None:
        paragraph = doc.add_paragraph()
        self._render_inline_content(paragraph, node.content or [])

    def _render_inline_content(self, paragraph, nodes: list[EditorNode]) -> None:
        """Render text nodes with marks into runs on a paragraph."""
        for node in nodes:
            if node.type == "text" and node.text:
                run = paragraph.add_run(node.text)
                self._apply_marks(run, node.marks or [])
            elif node.type == "hardBreak":
                run = paragraph.add_run()
                run.add_break()
            elif node.type == "docforgeUnresolved":
                text = node.attrs.get("originalText", "???") if node.attrs else "???"
                run = paragraph.add_run(text)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    def _apply_marks(self, run, marks: list[EditorMark]) -> None:
        """Apply TipTap marks to a python-docx Run."""
        for mark in marks:
            if mark.type == "bold":
                run.font.bold = True
            elif mark.type == "italic":
                run.font.italic = True
            elif mark.type == "underline":
                run.font.underline = True
            elif mark.type == "textStyle" and mark.attrs:
                if "fontSize" in mark.attrs:
                    pt_str = mark.attrs["fontSize"].replace("pt", "")
                    try:
                        run.font.size = Pt(float(pt_str))
                    except (ValueError, TypeError):
                        pass
                if "fontFamily" in mark.attrs:
                    run.font.name = mark.attrs["fontFamily"]
                if "color" in mark.attrs:
                    hex_color = mark.attrs["color"].lstrip("#")
                    try:
                        run.font.color.rgb = RGBColor.from_string(hex_color)
                    except (ValueError, TypeError):
                        pass
            # docforgeRendered marks are metadata-only — not applied to .docx formatting

    def _render_table(self, doc: Document, node: EditorNode) -> None:
        """Render a TipTap table node into a python-docx Table."""
        rows = node.content or []
        if not rows:
            return

        num_cols = len(rows[0].content or []) if rows[0].content else 0
        if num_cols == 0:
            return

        num_rows = len(rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row_node in enumerate(rows):
            cells = row_node.content or []
            for col_idx, cell_node in enumerate(cells):
                if col_idx >= num_cols:
                    break
                cell = table.rows[row_idx].cells[col_idx]
                # Clear default paragraph text
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                cell_content = cell_node.content or []
                for child_idx, child in enumerate(cell_content):
                    if child.type == "paragraph":
                        if child_idx == 0 and cell.paragraphs:
                            para = cell.paragraphs[0]
                        else:
                            para = cell.add_paragraph()
                        self._render_inline_content(para, child.content or [])

    def _render_unresolved_inline(self, doc: Document, node: EditorNode) -> None:
        """Render a standalone unresolved marker as a red text paragraph."""
        text = node.attrs.get("originalText", "???") if node.attrs else "???"
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
