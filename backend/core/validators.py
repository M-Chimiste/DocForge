"""Validation functions for mappings and generated output."""

from __future__ import annotations

from docx import Document

from core.data_loader import DataStore
from core.models import MappingEntry, TemplateAnalysis, ValidationIssue
from utils.red_text import is_red_run


def validate_mappings(
    analysis: TemplateAnalysis,
    mappings: list[MappingEntry],
    data_store: DataStore,
) -> list[ValidationIssue]:
    """Pre-generation validation: unmapped markers, missing sources, schema mismatches."""
    issues: list[ValidationIssue] = []
    mapping_by_marker = {m.marker_id: m for m in mappings}

    for marker in analysis.markers:
        mapping = mapping_by_marker.get(marker.id)

        if not mapping:
            # For table markers, check if there's a table-level mapping
            if marker.table_id and marker.table_id in mapping_by_marker:
                continue
            issues.append(
                ValidationIssue(
                    level="warning",
                    marker_id=marker.id,
                    message=f"Marker '{marker.text}' has no mapping configured",
                )
            )
            continue

        source = data_store.get(mapping.data_source)
        if source is None:
            issues.append(
                ValidationIssue(
                    level="error",
                    marker_id=marker.id,
                    message=f"Data source '{mapping.data_source}' not found",
                )
            )
            continue

        if mapping.field:
            df = data_store.get_dataframe(mapping.data_source, sheet=mapping.sheet)
            if df is not None and mapping.field not in df.columns:
                issues.append(
                    ValidationIssue(
                        level="error",
                        marker_id=marker.id,
                        message=f"Field '{mapping.field}' not found in '{mapping.data_source}'",
                    )
                )

    return issues


def validate_output(document: Document) -> list[ValidationIssue]:
    """Post-generation validation: check for unresolved red text."""
    issues: list[ValidationIssue] = []

    for para_idx, paragraph in enumerate(document.paragraphs):
        for run in paragraph.runs:
            if is_red_run(run._element):
                issues.append(
                    ValidationIssue(
                        level="warning",
                        message=f"Unresolved red text at paragraph {para_idx}: '{run.text[:50]}'",
                    )
                )

    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if is_red_run(run._element):
                            issues.append(
                                ValidationIssue(
                                    level="warning",
                                    message=(
                                        f"Unresolved red text in table {table_idx}, "
                                        f"row {row_idx}: '{run.text[:50]}'"
                                    ),
                                )
                            )

    return issues
