"""Template parser: extracts sections, markers, and tables from .docx templates."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from core.models import (
    MarkerType,
    Section,
    SkeletonTable,
    TemplateAnalysis,
    TemplateMarker,
)
from utils.docx_helpers import get_heading_level
from utils.red_text import classify_marker, is_red_run


def parse_template(template_path: Path) -> TemplateAnalysis:
    """Parse a .docx template and extract all sections, markers, and tables.

    Returns a TemplateAnalysis containing:
    - Sections detected from heading styles
    - Red text markers classified by the rule chain
    - Skeleton tables with header detection
    """
    doc = Document(str(template_path))

    sections: list[Section] = []
    markers: list[TemplateMarker] = []
    tables: list[SkeletonTable] = []

    current_section_id: str | None = None
    marker_counter = 0
    section_counter = 0
    table_counter = 0

    # Walk paragraphs for sections and inline red text
    for para_idx, paragraph in enumerate(doc.paragraphs):
        level = get_heading_level(paragraph)
        if level is not None:
            section_id = f"section-{section_counter}"
            section_counter += 1
            sections.append(
                Section(
                    id=section_id,
                    title=paragraph.text,
                    level=level,
                    paragraph_index=para_idx,
                )
            )
            current_section_id = section_id

        # Detect red text runs and group consecutive ones
        red_groups = _extract_red_groups(paragraph)
        for run_indices, text in red_groups:
            marker_id = f"marker-{marker_counter}"
            marker_counter += 1
            marker_type = classify_marker(text, in_table_data_row=False)
            marker = TemplateMarker(
                id=marker_id,
                text=text,
                marker_type=marker_type,
                section_id=current_section_id,
                paragraph_index=para_idx,
                run_indices=run_indices,
            )
            markers.append(marker)
            # Attach to current section
            if current_section_id:
                for s in sections:
                    if s.id == current_section_id:
                        s.markers.append(marker)
                        break

    # Walk tables for skeleton table detection
    for table in doc.tables:
        table_id = f"table-{table_counter}"
        table_counter += 1

        if not table.rows:
            continue

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        data_row_count = max(0, len(table.rows) - 1)

        # Check for red text in data rows (sample data markers)
        sample_markers: list[TemplateMarker] = []
        for row_idx in range(1, len(table.rows)):
            for cell in table.rows[row_idx].cells:
                for para in cell.paragraphs:
                    red_groups = _extract_red_groups(para)
                    for run_indices, text in red_groups:
                        mid = f"marker-{marker_counter}"
                        marker_counter += 1
                        m = TemplateMarker(
                            id=mid,
                            text=text,
                            marker_type=MarkerType.SAMPLE_DATA,
                            section_id=current_section_id,
                            paragraph_index=-1,
                            run_indices=run_indices,
                            table_id=table_id,
                            row_index=row_idx,
                        )
                        sample_markers.append(m)
                        markers.append(m)

        tables.append(
            SkeletonTable(
                id=table_id,
                section_id=current_section_id,
                paragraph_index=-1,
                headers=headers,
                row_count=data_row_count,
                sample_data_markers=sample_markers,
            )
        )

    return TemplateAnalysis(sections=sections, markers=markers, tables=tables)


def _extract_red_groups(paragraph) -> list[tuple[list[int], str]]:
    """Extract groups of consecutive red runs from a paragraph.

    Returns list of (run_indices, combined_text) tuples.
    Consecutive red runs are merged into a single group.
    """
    groups: list[tuple[list[int], str]] = []
    current_indices: list[int] = []
    current_texts: list[str] = []

    for run_idx, run in enumerate(paragraph.runs):
        if is_red_run(run._element):
            current_indices.append(run_idx)
            current_texts.append(run.text)
        else:
            if current_indices:
                groups.append((current_indices[:], "".join(current_texts)))
                current_indices.clear()
                current_texts.clear()

    # Don't forget trailing group
    if current_indices:
        groups.append((current_indices[:], "".join(current_texts)))

    return groups
