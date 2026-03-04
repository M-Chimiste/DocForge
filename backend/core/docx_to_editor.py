"""Convert a python-docx Document to TipTap-compatible EditorDocument JSON.

Walks the document body XML directly to get correct paragraph/table
interleaving, extracts formatting as TipTap marks, and annotates
rendered content with DocForge metadata.
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from core.editor_models import (
    EditorDocument,
    EditorDocumentMeta,
    EditorMark,
    EditorNode,
    MarkerEditorMeta,
)
from core.models import (
    AutoResolutionReport,
    GenerationReport,
    MappingEntry,
    RenderResult,
    TemplateAnalysis,
    TemplateMarker,
)
from utils.docx_helpers import get_heading_level
from utils.red_text import is_red_run

# Confidence threshold: auto-resolved markers below this are flagged in editor
LOW_CONFIDENCE_THRESHOLD = 0.8


class DocxToEditorConverter:
    """Convert a python-docx Document to TipTap-compatible EditorNode JSON."""

    def __init__(
        self,
        document: Document,
        generation_report: GenerationReport,
        analysis: TemplateAnalysis,
        mapping_snapshot: list[MappingEntry],
        auto_resolution_report: AutoResolutionReport | None = None,
        project_id: int = 0,
        run_id: int = 0,
    ):
        self._document = document
        self._report = generation_report
        self._analysis = analysis
        self._mappings = mapping_snapshot
        self._auto_report = auto_resolution_report
        self._project_id = project_id
        self._run_id = run_id

        # Build lookup indexes
        self._results_by_marker: dict[str, RenderResult] = {
            r.marker_id: r for r in generation_report.results
        }
        self._markers_by_para: dict[int, list[TemplateMarker]] = {}
        for marker in analysis.markers:
            self._markers_by_para.setdefault(marker.paragraph_index, []).append(marker)
        self._confidence_by_marker: dict[str, float] = {}
        if auto_resolution_report:
            for match in auto_resolution_report.matches:
                self._confidence_by_marker[match.marker_id] = match.confidence
        self._mappings_by_marker: dict[str, MappingEntry] = {
            m.marker_id: m for m in mapping_snapshot
        }
        self._sections_by_para: dict[int, str] = {
            s.paragraph_index: s.id for s in analysis.sections
        }

    def convert(self) -> EditorDocument:
        """Walk the document and produce the full EditorDocument."""
        body_nodes = self._walk_body_elements()
        root = EditorNode(type="doc", content=body_nodes)
        meta = EditorDocumentMeta(
            generation_run_id=self._run_id,
            project_id=self._project_id,
            template_analysis=self._analysis.model_dump(mode="json"),
            marker_metadata=self._build_marker_metadata(),
        )
        return EditorDocument(content=root, meta=meta)

    def _walk_body_elements(self) -> list[EditorNode]:
        """Walk w:body children in document order, handling both p and tbl."""
        nodes: list[EditorNode] = []
        para_idx = 0
        table_idx = 0
        for child in self._document.element.body:
            tag = child.tag
            if tag == qn("w:p"):
                if para_idx < len(self._document.paragraphs):
                    paragraph = self._document.paragraphs[para_idx]
                    node = self._convert_paragraph(paragraph, para_idx)
                    if node:
                        nodes.append(node)
                para_idx += 1
            elif tag == qn("w:tbl"):
                if table_idx < len(self._document.tables):
                    table = self._document.tables[table_idx]
                    node = self._convert_table(table, table_idx)
                    if node:
                        nodes.append(node)
                table_idx += 1
            # Skip w:sectPr and other elements
        return nodes

    def _convert_paragraph(self, paragraph: Paragraph, para_idx: int) -> EditorNode | None:
        """Convert a single paragraph to a heading or paragraph EditorNode."""
        heading_level = get_heading_level(paragraph)
        inline_nodes = self._convert_runs(paragraph, para_idx)

        if heading_level is not None:
            attrs: dict = {"level": heading_level}
            # Add section id for navigation if this paragraph starts a section
            section_id = self._sections_by_para.get(para_idx)
            if section_id:
                attrs["id"] = f"section-{section_id}"
            return EditorNode(
                type="heading",
                attrs=attrs,
                content=inline_nodes if inline_nodes else None,
            )

        return EditorNode(
            type="paragraph",
            content=inline_nodes if inline_nodes else None,
        )

    def _convert_runs(self, paragraph: Paragraph, para_idx: int) -> list[EditorNode]:
        """Convert paragraph runs to text nodes with marks."""
        nodes: list[EditorNode] = []
        runs = paragraph.runs
        markers_here = self._markers_by_para.get(para_idx, [])

        # Build a set of (run_index -> marker) for quick lookup
        run_to_marker: dict[int, TemplateMarker] = {}
        for marker in markers_here:
            if marker.table_id is None:  # Skip table markers
                for ri in marker.run_indices:
                    run_to_marker[ri] = marker

        for run_idx, run in enumerate(runs):
            text = run.text
            if not text:
                continue

            marks = self._extract_marks(run)
            marker = run_to_marker.get(run_idx)

            if marker:
                result = self._results_by_marker.get(marker.id)
                if result and result.success:
                    # Rendered content — add docforgeRendered mark
                    rendered_by = result.rendered_by or "unknown"
                    confidence = self._confidence_by_marker.get(marker.id)

                    # Only flag if LLM-rendered or low-confidence
                    should_flag = rendered_by == "llm" or (
                        confidence is not None and confidence < LOW_CONFIDENCE_THRESHOLD
                    )
                    if should_flag:
                        marks.append(
                            EditorMark(
                                type="docforgeRendered",
                                attrs={
                                    "markerId": marker.id,
                                    "renderType": rendered_by,
                                    "confidence": confidence,
                                },
                            )
                        )
                elif is_red_run(run._element):
                    # Unresolved marker — still red text
                    nodes.append(
                        EditorNode(
                            type="docforgeUnresolved",
                            attrs={
                                "markerId": marker.id,
                                "originalText": marker.text,
                            },
                        )
                    )
                    continue
            elif is_red_run(run._element):
                # Red run not matched to a known marker — treat as unresolved
                nodes.append(
                    EditorNode(
                        type="docforgeUnresolved",
                        attrs={
                            "markerId": "",
                            "originalText": text,
                        },
                    )
                )
                continue

            nodes.append(
                EditorNode(
                    type="text",
                    text=text,
                    marks=marks if marks else None,
                )
            )

        return nodes

    def _extract_marks(self, run: Run) -> list[EditorMark]:
        """Extract TipTap formatting marks from a python-docx Run."""
        marks: list[EditorMark] = []
        font = run.font

        if font.bold:
            marks.append(EditorMark(type="bold"))
        if font.italic:
            marks.append(EditorMark(type="italic"))
        if font.underline:
            marks.append(EditorMark(type="underline"))

        # Collect textStyle attrs
        style_attrs: dict[str, str] = {}
        if font.size:
            # python-docx stores size as EMU, .pt gives points
            pt_val = font.size.pt
            style_attrs["fontSize"] = f"{pt_val}pt"
        if font.name:
            style_attrs["fontFamily"] = font.name
        if font.color and font.color.rgb:
            style_attrs["color"] = f"#{font.color.rgb}"

        if style_attrs:
            marks.append(EditorMark(type="textStyle", attrs=style_attrs))

        return marks

    def _convert_table(self, table: DocxTable, table_idx: int) -> EditorNode:
        """Convert a python-docx Table to TipTap table structure."""
        row_nodes: list[EditorNode] = []

        for row_idx, row in enumerate(table.rows):
            cell_nodes: list[EditorNode] = []
            for cell in row.cells:
                # Convert cell paragraphs
                cell_content: list[EditorNode] = []
                for para in cell.paragraphs:
                    inline_nodes = self._convert_cell_runs(para)
                    cell_content.append(
                        EditorNode(
                            type="paragraph",
                            content=inline_nodes if inline_nodes else None,
                        )
                    )

                # First row uses tableHeader, rest use tableCell
                cell_type = "tableHeader" if row_idx == 0 else "tableCell"
                cell_nodes.append(
                    EditorNode(
                        type=cell_type,
                        content=cell_content if cell_content else None,
                    )
                )

            row_nodes.append(
                EditorNode(type="tableRow", content=cell_nodes if cell_nodes else None)
            )

        return EditorNode(
            type="table",
            attrs={"docforgeTableId": f"table-{table_idx}"},
            content=row_nodes if row_nodes else None,
        )

    def _convert_cell_runs(self, paragraph: Paragraph) -> list[EditorNode]:
        """Convert runs in a table cell paragraph to text nodes."""
        nodes: list[EditorNode] = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            marks = self._extract_marks(run)
            nodes.append(
                EditorNode(
                    type="text",
                    text=text,
                    marks=marks if marks else None,
                )
            )

        # If no runs but paragraph has text (cell.text setter creates no runs)
        if not nodes and paragraph.text:
            nodes.append(EditorNode(type="text", text=paragraph.text))

        return nodes

    def _build_marker_metadata(self) -> dict[str, MarkerEditorMeta]:
        """Build per-marker metadata from generation report and analysis."""
        metadata: dict[str, MarkerEditorMeta] = {}

        for marker in self._analysis.markers:
            result = self._results_by_marker.get(marker.id)
            confidence = self._confidence_by_marker.get(marker.id)
            mapping = self._mappings_by_marker.get(marker.id)

            if result and result.success:
                rendered_by = result.rendered_by or "unknown"
            else:
                rendered_by = "unresolved"

            llm_model = None
            llm_prompt_tokens = None
            if result and result.llm_usage:
                llm_model = result.llm_usage.get("model")
                llm_prompt_tokens = result.llm_usage.get("prompt_tokens")

            metadata[marker.id] = MarkerEditorMeta(
                marker_id=marker.id,
                marker_type=marker.marker_type.value,
                original_text=marker.text,
                rendered_by=rendered_by,
                confidence=confidence,
                llm_model=llm_model,
                llm_prompt_tokens=llm_prompt_tokens,
                section_id=marker.section_id,
                mapping_snapshot=mapping.model_dump(mode="json") if mapping else None,
            )

        return metadata
