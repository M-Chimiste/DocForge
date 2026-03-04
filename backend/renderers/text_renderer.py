"""TextRenderer: inject multi-paragraph text content into designated sections."""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.data_loader import DataStore
from core.models import MappingEntry, MarkerType, RenderResult, TemplateMarker
from renderers.base import BaseRenderer
from utils.docx_helpers import copy_run_format, find_adjacent_non_red_run
from utils.formatting import clear_red_color


class TextRenderer(BaseRenderer):
    """Handles LLM_PROMPT markers mapped to text data sources.

    When an LLM_PROMPT marker is mapped to a data source with text_content,
    the TextRenderer injects that text directly. This allows text file content
    to fill template sections without LLM processing.
    """

    def can_handle(self, marker: TemplateMarker) -> bool:
        return marker.marker_type == MarkerType.LLM_PROMPT

    def render(
        self,
        marker: TemplateMarker,
        data: DataStore,
        document: Document,
        mapping: MappingEntry,
    ) -> RenderResult:
        # Get text content from data source
        text = data.get_text(mapping.data_source)
        if text is None:
            # Try getting from DataFrame's "content" column
            df = data.get_dataframe(mapping.data_source, sheet=mapping.sheet)
            if df is not None and "content" in df.columns:
                text = str(df.iloc[0]["content"])
            else:
                return RenderResult(
                    marker_id=marker.id,
                    success=False,
                    error=f"No text content found in {mapping.data_source}",
                )

        if marker.paragraph_index < 0 or marker.paragraph_index >= len(document.paragraphs):
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error=f"Paragraph index {marker.paragraph_index} out of range",
            )

        paragraph = document.paragraphs[marker.paragraph_index]
        runs = paragraph.runs

        # Get formatting from adjacent non-red run
        format_source = find_adjacent_non_red_run(paragraph, marker.run_indices)

        # Split text into paragraphs
        text_parts = text.strip().split("\n")

        # Replace the red text runs
        for i, run_idx in enumerate(marker.run_indices):
            if run_idx >= len(runs):
                continue
            run = runs[run_idx]
            if i == 0:
                run.text = text_parts[0] if text_parts else ""
            else:
                run.text = ""
            clear_red_color(run)
            if format_source:
                copy_run_format(format_source, run)

        # Add subsequent paragraphs after the current one
        if len(text_parts) > 1:
            current_para = paragraph
            for extra_text in text_parts[1:]:
                if extra_text.strip():
                    new_para = _insert_paragraph_after(current_para, extra_text)
                    if format_source and new_para.runs:
                        copy_run_format(format_source, new_para.runs[0])
                    current_para = new_para

        return RenderResult(marker_id=marker.id, success=True)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Insert a new paragraph after the given paragraph."""
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para
