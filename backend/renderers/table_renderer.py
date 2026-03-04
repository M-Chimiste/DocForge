"""Table renderer: populates skeleton tables from DataFrames."""

from __future__ import annotations

from docx import Document
from docx.table import Table

from core.data_loader import DataStore
from core.models import MappingEntry, MarkerType, RenderResult, TemplateMarker
from renderers.base import BaseRenderer


class TableRenderer(BaseRenderer):
    def can_handle(self, marker: TemplateMarker) -> bool:
        return marker.marker_type == MarkerType.SAMPLE_DATA

    def render(
        self,
        marker: TemplateMarker,
        data: DataStore,
        document: Document,
        mapping: MappingEntry,
    ) -> RenderResult:
        if not marker.table_id:
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error="Marker has no table_id",
            )

        df = data.get_dataframe(mapping.data_source, sheet=mapping.sheet)
        if df is None:
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error=f"No data found for {mapping.data_source}",
            )

        # Find the table in the document by matching table_id index
        table_index = int(marker.table_id.split("-")[1])
        if table_index >= len(document.tables):
            return RenderResult(
                marker_id=marker.id,
                success=False,
                error=f"Table index {table_index} out of range",
            )

        table = document.tables[table_index]
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Remove existing data rows (keep header)
        _remove_data_rows(table)

        # Map DataFrame columns to table headers
        col_mapping = _map_columns(headers, list(df.columns))

        # Add rows from DataFrame
        for _, row_data in df.iterrows():
            new_row = table.add_row()
            for col_idx, header in enumerate(headers):
                mapped_col = col_mapping.get(header)
                if mapped_col and mapped_col in row_data.index:
                    value = row_data[mapped_col]
                    new_row.cells[col_idx].text = str(value) if value is not None else ""

        return RenderResult(marker_id=marker.id, success=True)


def render_table_direct(
    table_id: str,
    data: DataStore,
    document: Document,
    mapping: MappingEntry,
) -> RenderResult:
    """Render a table directly without a marker (for skeleton tables with no sample data)."""
    df = data.get_dataframe(mapping.data_source, sheet=mapping.sheet)
    if df is None:
        return RenderResult(
            marker_id=table_id,
            success=False,
            error=f"No data found for {mapping.data_source}",
        )

    table_index = int(table_id.split("-")[1])
    if table_index >= len(document.tables):
        return RenderResult(
            marker_id=table_id,
            success=False,
            error=f"Table index {table_index} out of range",
        )

    table = document.tables[table_index]
    headers = [cell.text.strip() for cell in table.rows[0].cells]

    col_mapping = _map_columns(headers, list(df.columns))

    for _, row_data in df.iterrows():
        new_row = table.add_row()
        for col_idx, header in enumerate(headers):
            mapped_col = col_mapping.get(header)
            if mapped_col and mapped_col in row_data.index:
                value = row_data[mapped_col]
                new_row.cells[col_idx].text = str(value) if value is not None else ""

    return RenderResult(marker_id=table_id, success=True)


def _remove_data_rows(table: Table) -> None:
    """Remove all rows except the header row."""
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)


def _map_columns(table_headers: list[str], df_columns: list[str]) -> dict[str, str]:
    """Map table headers to DataFrame columns by exact or case-insensitive match."""
    mapping: dict[str, str] = {}
    df_lower = {col.lower(): col for col in df_columns}

    for header in table_headers:
        if header in df_columns:
            mapping[header] = header
        elif header.lower() in df_lower:
            mapping[header] = df_lower[header.lower()]

    return mapping
